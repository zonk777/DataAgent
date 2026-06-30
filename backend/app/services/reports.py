from __future__ import annotations

import html
import json
import math
import os
import tempfile
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image as PdfImage
from reportlab.platypus import LongTable, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from ..db import connect


FONT_CANDIDATES = [
    Path("C:/Windows/Fonts/Deng.ttf"),
    Path("C:/Windows/Fonts/msyh.ttc"),
    Path("C:/Windows/Fonts/simhei.ttf"),
    Path("C:/Windows/Fonts/simsun.ttc"),
]


@dataclass
class ReportData:
    session: dict[str, Any]
    assistant_message: dict[str, Any]
    question: str
    payload: dict[str, Any]


def load_report_data(session_id: str) -> ReportData:
    with connect() as conn:
        session = conn.execute("SELECT * FROM sessions WHERE id = ?", (session_id,)).fetchone()
        if not session:
            raise LookupError("会话不存在")
        assistant = conn.execute(
            """SELECT id, content, payload, created_at
               FROM messages
               WHERE session_id = ? AND role = 'assistant' AND payload IS NOT NULL
               ORDER BY id DESC LIMIT 1""",
            (session_id,),
        ).fetchone()
        if not assistant:
            raise LookupError("该会话还没有可导出的分析结果")
        user = conn.execute(
            """SELECT content
               FROM messages
               WHERE session_id = ? AND role = 'user' AND id < ?
               ORDER BY id DESC LIMIT 1""",
            (session_id, assistant["id"]),
        ).fetchone()
    payload = json.loads(assistant["payload"])
    return ReportData(
        session=dict(session),
        assistant_message=dict(assistant),
        question=user["content"] if user else payload.get("effective_question", ""),
        payload=payload,
    )


def _safe_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value)


def _report_title(data: ReportData) -> str:
    return _safe_text(data.payload.get("chart", {}).get("title") or data.session.get("title") or "数据智能体分析报告")


def build_html_report(data: ReportData) -> str:
    payload = data.payload
    insights = "".join(f"<li>{html.escape(str(item))}</li>" for item in payload.get("insights", []))
    references = "".join(
        f"<li><strong>{html.escape(str(item.get('title', '')))}</strong>：{html.escape(str(item.get('content', '')))}</li>"
        for item in payload.get("knowledge_refs", [])
    )
    table = ""
    if payload.get("rows"):
        columns = payload.get("columns", [])
        table_head = "".join(f"<th>{html.escape(str(column))}</th>" for column in columns)
        table_rows = "".join(
            "<tr>" + "".join(f"<td>{html.escape(str(item.get(column, '')))}</td>" for column in columns) + "</tr>"
            for item in payload["rows"]
        )
        table = f"<h2>查询结果</h2><table><thead><tr>{table_head}</tr></thead><tbody>{table_rows}</tbody></table>"
    sql = f"<h2>执行 SQL</h2><pre>{html.escape(payload['sql'])}</pre>" if payload.get("sql") else ""
    refs = f"<h2>知识依据</h2><ul>{references}</ul>" if references else ""
    return f"""<!doctype html><html lang="zh-CN"><head><meta charset="utf-8"><title>数据智能体分析报告</title>
    <style>body{{font-family:Arial,'Microsoft YaHei',sans-serif;max-width:980px;margin:48px auto;color:#16324f;line-height:1.7}}
    h1{{color:#087ea4}} .meta{{color:#667085}} table{{width:100%;border-collapse:collapse;margin:24px 0}}
    th,td{{padding:10px 12px;border:1px solid #dbe5ee;text-align:left}} th{{background:#edf8fc}} pre{{background:#f4f7fa;padding:16px;white-space:pre-wrap}}</style></head>
    <body><h1>{html.escape(_report_title(data))}</h1>
    <p class="meta">会话编号：{html.escape(data.session["id"])} · 类型：{html.escape(payload.get("intent", ""))} · 生成时间：{html.escape(data.assistant_message.get("created_at", ""))}</p>
    <p><strong>分析问题：</strong>{html.escape(data.question)}</p>
    <h2>回答与发现</h2><ul>{insights}</ul>{table}{sql}{refs}</body></html>"""


def _find_font_path() -> Path | None:
    return next((path for path in FONT_CANDIDATES if path.exists()), None)


def _image_font(size: int, bold: bool = False) -> ImageFont.ImageFont:
    path = Path("C:/Windows/Fonts/Dengb.ttf") if bold and Path("C:/Windows/Fonts/Dengb.ttf").exists() else _find_font_path()
    if path:
        return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def _float_value(value: Any) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return 0.0


def _series_key(row: dict[str, Any], series_fields: list[str]) -> str:
    if not series_fields:
        return "结果"
    return " / ".join(_safe_text(row.get(field) or "未分类") for field in series_fields)


def build_chart_image(payload: dict[str, Any]) -> BytesIO | None:
    chart = payload.get("chart") or {}
    rows = payload.get("rows") or []
    chart_type = chart.get("type")
    x_field = chart.get("x_field")
    y_field = chart.get("y_field")
    if not rows or not x_field or not y_field or chart_type == "none":
        return None

    width, height = 1100, 520
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = _image_font(28, True)
    label_font = _image_font(18)
    small_font = _image_font(15)
    axis_color = (76, 100, 132)
    grid_color = (224, 234, 242)
    colorset = [(22, 119, 255), (16, 183, 196), (91, 117, 231), (245, 158, 11), (16, 163, 127), (239, 71, 111)]
    draw.text((42, 26), _safe_text(chart.get("title") or "分析图表"), font=title_font, fill=(20, 44, 76))
    plot = (78, 95, 1040, 420)

    series_fields = chart.get("series_fields") or ([chart["series_field"]] if chart.get("series_field") else [])
    grouped: dict[str, dict[str, float]] = {}
    x_values: list[str] = []
    for row in rows:
        x = _safe_text(row.get(x_field))
        s = _series_key(row, series_fields)
        if x not in x_values:
            x_values.append(x)
        grouped.setdefault(s, {})[x] = grouped.setdefault(s, {}).get(x, 0) + _float_value(row.get(y_field))
    x_values = x_values[:28]

    if chart_type == "pie":
        totals = [(x, sum(group.get(x, 0) for group in grouped.values())) for x in x_values[:10]]
        total = sum(value for _, value in totals) or 1
        bbox = (165, 120, 475, 430)
        start = 0.0
        for idx, (label, value) in enumerate(totals):
            extent = value / total * 360
            draw.pieslice(bbox, start, start + extent, fill=colorset[idx % len(colorset)], outline="white", width=2)
            start += extent
            y = 125 + idx * 30
            draw.rectangle((565, y + 3, 585, y + 23), fill=colorset[idx % len(colorset)])
            draw.text((595, y), f"{label}  {value:,.2f}", font=label_font, fill=(40, 58, 82))
    else:
        values = [value for group in grouped.values() for value in group.values()]
        max_value = max(values) if values else 1
        min_value = min(0, min(values) if values else 0)
        span = max(max_value - min_value, 1)
        x1, y1, x2, y2 = plot
        for idx in range(6):
            y = y2 - int((y2 - y1) * idx / 5)
            draw.line((x1, y, x2, y), fill=grid_color, width=1)
            label = min_value + span * idx / 5
            draw.text((10, y - 10), f"{label:,.0f}", font=small_font, fill=(112, 129, 151))
        draw.line((x1, y2, x2, y2), fill=axis_color, width=2)
        draw.line((x1, y1, x1, y2), fill=axis_color, width=2)
        if chart_type == "line":
            for s_idx, (series, data) in enumerate(list(grouped.items())[:6]):
                pts = []
                for idx, x in enumerate(x_values):
                    px = x1 + int((x2 - x1) * idx / max(len(x_values) - 1, 1))
                    py = y2 - int((data.get(x, 0) - min_value) / span * (y2 - y1))
                    pts.append((px, py))
                if len(pts) > 1:
                    draw.line(pts, fill=colorset[s_idx % len(colorset)], width=4)
                for px, py in pts:
                    draw.ellipse((px - 4, py - 4, px + 4, py + 4), fill=colorset[s_idx % len(colorset)])
        else:
            totals = [(x, sum(group.get(x, 0) for group in grouped.values())) for x in x_values]
            bar_w = max(10, int((x2 - x1) / max(len(totals), 1) * 0.62))
            for idx, (x, value) in enumerate(totals):
                cx = x1 + int((x2 - x1) * (idx + 0.5) / max(len(totals), 1))
                top = y2 - int((value - min_value) / span * (y2 - y1))
                draw.rounded_rectangle((cx - bar_w // 2, top, cx + bar_w // 2, y2), radius=5, fill=colorset[idx % len(colorset)])
        label_step = max(1, math.ceil(len(x_values) / 8))
        for idx, x in enumerate(x_values):
            if idx % label_step:
                continue
            px = x1 + int((x2 - x1) * idx / max(len(x_values) - 1, 1))
            draw.text((px - 38, y2 + 12), x[:10], font=small_font, fill=(93, 111, 132))
        for s_idx, series in enumerate(list(grouped.keys())[:6]):
            lx = 82 + s_idx * 150
            draw.rectangle((lx, 452, lx + 18, 470), fill=colorset[s_idx % len(colorset)])
            draw.text((lx + 26, 448), series[:10], font=small_font, fill=(70, 88, 112))

    out = BytesIO()
    img.save(out, format="PNG")
    out.seek(0)
    return out


def _set_run_font(run, size: int | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = color


def _add_docx_paragraph(doc: Document, text: str, size: int = 10, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold)


def _add_docx_table(doc: Document, headers: list[str], rows: list[list[Any]], max_rows: int = 80) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = _safe_text(header)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                _set_run_font(run, size=9, bold=True, color=RGBColor(11, 37, 69))
    for row in rows[:max_rows]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = _safe_text(value)
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    _set_run_font(run, size=8)
    if len(rows) > max_rows:
        _add_docx_paragraph(doc, f"注：结果共 {len(rows)} 行，Word 报告仅展示前 {max_rows} 行。", size=9)


def build_docx_report(data: ReportData) -> bytes:
    payload = data.payload
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("数据智能体分析报告")
    _set_run_font(run, size=22, bold=True, color=RGBColor(11, 37, 69))
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(_report_title(data))
    _set_run_font(run, size=13, bold=True, color=RGBColor(46, 116, 181))

    doc.add_heading("一、报告概况", level=1)
    _add_docx_table(
        doc,
        ["项目", "内容"],
        [
            ["会话编号", data.session["id"]],
            ["分析问题", data.question],
            ["分析类型", payload.get("intent", "")],
            ["执行模式", payload.get("execution_mode", "")],
            ["生成时间", data.assistant_message.get("created_at", "")],
        ],
    )

    doc.add_heading("二、回答与关键发现", level=1)
    for idx, insight in enumerate(payload.get("insights") or [], 1):
        _add_docx_paragraph(doc, f"{idx}. {insight}", size=10)

    chart_image = build_chart_image(payload)
    if chart_image:
        doc.add_heading("三、结果图表", level=1)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(chart_image, width=Inches(6.4))

    if payload.get("rows"):
        doc.add_heading("四、查询结果", level=1)
        columns = payload.get("columns", [])
        _add_docx_table(doc, columns, [[row.get(column, "") for column in columns] for row in payload.get("rows", [])])

    if payload.get("sql"):
        doc.add_heading("五、执行 SQL", level=1)
        _add_docx_paragraph(doc, payload["sql"], size=9)

    if payload.get("knowledge_refs"):
        doc.add_heading("六、知识依据", level=1)
        _add_docx_table(
            doc,
            ["标题", "类别", "内容"],
            [[item.get("title", ""), item.get("category", ""), item.get("content", "")] for item in payload["knowledge_refs"]],
            max_rows=20,
        )

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _register_pdf_font() -> str:
    font_name = "DataAgentCN"
    if font_name in pdfmetrics.getRegisteredFontNames():
        return font_name
    font_path = _find_font_path()
    if font_path:
        pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
        return font_name
    return "Helvetica"


def _pdf_styles() -> dict[str, ParagraphStyle]:
    font_name = _register_pdf_font()
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle("title_cn", parent=base["Title"], fontName=font_name, fontSize=22, leading=28, alignment=TA_CENTER, textColor=colors.HexColor("#0B2545"), spaceAfter=12),
        "subtitle": ParagraphStyle("subtitle_cn", parent=base["Normal"], fontName=font_name, fontSize=12, leading=17, alignment=TA_CENTER, textColor=colors.HexColor("#2E74B5"), spaceAfter=16),
        "h1": ParagraphStyle("h1_cn", parent=base["Heading1"], fontName=font_name, fontSize=14, leading=19, textColor=colors.HexColor("#2E74B5"), spaceBefore=10, spaceAfter=8),
        "body": ParagraphStyle("body_cn", parent=base["Normal"], fontName=font_name, fontSize=9.5, leading=15, alignment=TA_LEFT, spaceAfter=5),
        "small": ParagraphStyle("small_cn", parent=base["Normal"], fontName=font_name, fontSize=8, leading=12, alignment=TA_LEFT),
    }


def _p(text: Any, style: ParagraphStyle) -> Paragraph:
    return Paragraph(html.escape(_safe_text(text)).replace("\n", "<br/>"), style)


def _pdf_table(headers: list[str], rows: list[list[Any]], styles: dict[str, ParagraphStyle], max_rows: int = 80) -> Table:
    data = [[_p(header, styles["small"]) for header in headers]]
    for row in rows[:max_rows]:
        data.append([_p(value, styles["small"]) for value in row])
    table = LongTable(data, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F0F6F9")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0B2545")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D9E5EE")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def build_pdf_report(data: ReportData) -> bytes:
    payload = data.payload
    out = BytesIO()
    doc = SimpleDocTemplate(out, pagesize=landscape(A4), leftMargin=14 * mm, rightMargin=14 * mm, topMargin=14 * mm, bottomMargin=14 * mm)
    styles = _pdf_styles()
    story: list[Any] = [
        Paragraph("数据智能体分析报告", styles["title"]),
        Paragraph(_report_title(data), styles["subtitle"]),
        Paragraph("一、报告概况", styles["h1"]),
        _pdf_table(
            ["项目", "内容"],
            [
                ["会话编号", data.session["id"]],
                ["分析问题", data.question],
                ["分析类型", payload.get("intent", "")],
                ["执行模式", payload.get("execution_mode", "")],
                ["生成时间", data.assistant_message.get("created_at", "")],
            ],
            styles,
        ),
        Spacer(1, 8),
        Paragraph("二、回答与关键发现", styles["h1"]),
    ]
    for idx, insight in enumerate(payload.get("insights") or [], 1):
        story.append(Paragraph(f"{idx}. {html.escape(_safe_text(insight))}", styles["body"]))

    chart_image = build_chart_image(payload)
    if chart_image:
        story.extend([Spacer(1, 8), Paragraph("三、结果图表", styles["h1"])])
        tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        tmp.write(chart_image.getvalue())
        tmp.close()
        story.append(PdfImage(tmp.name, width=245 * mm, height=116 * mm))
        temp_chart_path = tmp.name
    else:
        temp_chart_path = None

    if payload.get("rows"):
        story.extend([PageBreak(), Paragraph("四、查询结果", styles["h1"])])
        columns = payload.get("columns", [])
        story.append(_pdf_table(columns, [[row.get(column, "") for column in columns] for row in payload.get("rows", [])], styles))
    if payload.get("sql"):
        story.extend([Spacer(1, 8), Paragraph("五、执行 SQL", styles["h1"]), Paragraph(payload["sql"], styles["small"])])
    if payload.get("knowledge_refs"):
        story.extend([Spacer(1, 8), Paragraph("六、知识依据", styles["h1"])])
        story.append(
            _pdf_table(
                ["标题", "类别", "内容"],
                [[item.get("title", ""), item.get("category", ""), item.get("content", "")] for item in payload["knowledge_refs"]],
                styles,
                max_rows=20,
            )
        )
    try:
        doc.build(story)
    finally:
        if temp_chart_path:
            try:
                os.remove(temp_chart_path)
            except OSError:
                pass
    return out.getvalue()
