from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT1 = ROOT / "docs" / "《数据智能体服务系统》概要设计-1.docx"
OUT2 = ROOT / "docs" / "《数据智能体服务系统》概要设计-2.docx"

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
GRAY = RGBColor(90, 90, 90)
HEADER_FILL = "D9EAF7"
LIGHT_FILL = "F2F4F7"
NOTE_FILL = "F7FBFD"


def set_run_font(run, size=None, bold=None, color=None, name="Calibri", east_asia="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def pformat(p, before=0, after=6, line=1.15, align=None):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align


def add_para(doc, text="", size=11, bold=False, color=None, style=None, before=0, after=6, align=None):
    p = doc.add_paragraph(style=style)
    pformat(p, before=before, after=after, align=align)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    pformat(p, after=4)
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    pformat(p, after=4)
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        item = tc_mar.find(qn(f"w:{name}"))
        if item is None:
            item = OxmlElement(f"w:{name}")
            tc_mar.append(item)
        item.set(qn("w:w"), str(value))
        item.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[i] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[i]))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def cell_text(cell, text, bold=False, size=9.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    pformat(p, after=0, line=1.12, align=align)
    parts = str(text).split("\n")
    for idx, part in enumerate(parts):
        if idx:
            p.add_run().add_break()
        r = p.add_run(part)
        set_run_font(r, size=size, bold=bold)


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_FILL)
        cell_text(cell, h, bold=True, size=font_size, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value, size=font_size)
    doc.add_paragraph()
    return table


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, NOTE_FILL)
    p = cell.paragraphs[0]
    pformat(p, after=3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=NAVY)
    p2 = cell.add_paragraph()
    pformat(p2, after=0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10)
    doc.add_paragraph()


def configure_doc(doc, subtitle):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 14, 8),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 11.5, DARK, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    pformat(header, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    r = header.add_run(f"数据智能体服务系统 | {subtitle}")
    set_run_font(r, size=8.5, color=GRAY)


def add_cover(doc, subtitle):
    for _ in range(4):
        doc.add_paragraph()
    p = add_para(doc, "《数据智能体服务系统》", size=26, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    p = add_para(doc, subtitle, size=20, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=28)
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["文档版本", "V1.0"],
            ["适用阶段", "概要设计"],
            ["项目周期", "3 周"],
            ["编写日期", "2026 年 6 月 24 日"],
            ["编写单位", "数据智能体项目组"],
        ],
        [2200, 7160],
        font_size=10.5,
    )
    add_note(
        doc,
        "文档说明",
        "本文档依据项目需求说明书和 D03 概要设计任务要求编写，用于承接需求分析并指导后续详细设计、编码实现和测试验收。",
    )
    doc.add_page_break()


def add_common_intro(doc):
    doc.add_heading("第一部分 引言", level=1)
    doc.add_heading("一、说明", level=2)
    add_para(
        doc,
        "本文档描述数据智能体服务系统的概要设计方案，包括项目概述、设计约束、功能结构、实体结构、总体架构、界面设计、算法设计和部署运行要求。概要设计承接需求说明，作为详细设计和编码实现的依据。",
    )
    doc.add_heading("二、定义", level=2)
    add_table(
        doc,
        ["术语", "说明"],
        [
            ["数据智能体", "能够根据自然语言问题自动完成数据理解、查询规划、分析执行、图表展示和结果解释的软件服务。"],
            ["RAG", "检索增强生成，通过知识库检索补充大模型回答依据。"],
            ["Embedding", "将文本转换为向量，用于语义检索。"],
            ["Qdrant Local", "本项目使用的轻量级本地向量数据库，用于保存知识片段向量。"],
            ["只读 SQL", "仅允许 SELECT/WITH 查询，不允许修改或删除数据的 SQL。"],
            ["多轮会话", "系统保存历史上下文，支持用户基于上一轮结果继续追问。"],
        ],
        [1800, 7560],
    )
    doc.add_heading("三、参考资料", level=2)
    for item in [
        "《数据智能体服务系统》项目需求说明.docx",
        "D03_概要设计-1.pdf：功能结构与实体结构设计要求",
        "D03_概要设计-2.pdf：UI、技术架构与总体设计要求",
        "项目源码：frontend、backend、docs、test-data、scripts 目录",
    ]:
        add_bullet(doc, item)


def add_project_overview(doc):
    doc.add_heading("第二部分 项目概述", level=1)
    doc.add_heading("一、项目描述", level=2)
    add_para(
        doc,
        "数据智能体服务系统面向企业经营数据分析场景，提供自然语言问数、数据源接入、业务知识库问答、多轮追问、图表展示、历史会话和报告导出等能力。系统采用 Vue 前端、FastAPI 后端、SQLite 结构化存储和 Qdrant Local 向量存储，形成轻量化、可本地运行的智能分析平台。",
    )
    doc.add_heading("二、项目功能描述", level=2)
    add_table(
        doc,
        ["功能域", "功能描述", "主要使用者"],
        [
            ["登录与账户管理", "网站打开后首先进入账号密码登录页；初始管理员可新增普通管理员，普通管理员无权新增其他管理员。", "系统管理员"],
            ["数据源管理", "支持 CSV/Excel 上传、数据入库、字段类型识别、样例值和空值率统计。", "数据分析人员"],
            ["智能分析", "用户输入中文问题，系统识别指标、时间、维度并生成只读 SQL，返回图表、表格和关键发现。", "管理者、运营人员"],
            ["多轮追问", "系统保存上下文，支持“只看华东”“按产品类别拆分”“改成柱状图”等追问。", "所有用户"],
            ["业务知识库", "维护指标口径、字段说明和业务规则，支持 Embedding + Qdrant Local 语义检索。", "系统管理员、分析人员"],
            ["历史会话", "保存用户消息、助手回答和分析 payload，支持恢复查看和删除。", "所有用户"],
            ["报告导出", "按会话导出 HTML 报告，包含结论、数据结果、SQL 和知识依据。", "分析人员"],
        ],
        [1900, 5260, 2200],
    )


def add_constraints(doc):
    doc.add_heading("第三部分 设计约束", level=1)
    add_table(
        doc,
        ["约束类别", "设计约束"],
        [
            ["周期约束", "项目周期为三周，优先完成可演示、可运行、可测试的 MVP，企业级扩展列入后续计划。"],
            ["技术约束", "前端采用 Vue 3 + TypeScript + Vite + ECharts；后端采用 FastAPI + SQLite；向量库采用 Qdrant Local。"],
            ["安全约束", "API Key 仅由后端 .env 读取；前端不展示密钥；业务接口必须登录后访问；SQL 只允许只读查询。"],
            ["部署约束", "本期以 Windows 本地部署为主，前端端口 5173，后端端口 8000。"],
            ["数据约束", "当前支持 CSV/Excel 和内置演示数据，后续可扩展 MySQL、PostgreSQL 等数据源。"],
            ["模型约束", "LLM 与 Embedding 通过 OpenAI 兼容接口接入，服务不可用时采用本地规则或关键词检索降级。"],
        ],
        [1900, 7460],
    )


def add_solution_design(doc):
    doc.add_heading("第四部分 数据智能体服务系统方案设计", level=1)
    doc.add_heading("一、安全设计", level=2)
    add_table(
        doc,
        ["安全点", "设计方案"],
        [
            ["登录认证", "系统打开后首先显示登录页；后端通过 HttpOnly Cookie 维护登录态，未登录访问业务接口返回 401。"],
            ["管理员权限", "系统初始化唯一初始管理员 liuze；只有初始管理员可以新增普通管理员；普通管理员新增管理员返回 403。"],
            ["密码保护", "管理员密码使用 PBKDF2-SHA256 加盐哈希保存，不在数据库中保存明文密码。"],
            ["SQL 安全", "只允许 SELECT/WITH 查询，拦截 DELETE、UPDATE、DROP、INSERT、ALTER 等危险语句。"],
            ["密钥保护", "LLM_API_KEY 与 EMBEDDING_API_KEY 保存在后端 .env，前端仅展示配置状态。"],
            ["审计记录", "系统记录分析操作、会话和导出动作，为后续审计与质量追踪提供依据。"],
        ],
        [1900, 7460],
    )
    doc.add_heading("二、数据智能体服务系统相关业务流程", level=2)
    add_note(
        doc,
        "自然语言分析流程",
        "用户登录 -> 选择数据集 -> 输入问题 -> 智能体识别意图 -> 解析指标/时间/维度 -> 生成只读 SQL -> 安全校验 -> 执行查询 -> 生成 chart spec -> 前端渲染图表 -> 保存会话历史。",
    )
    add_note(
        doc,
        "知识库问答流程",
        "用户提出指标口径问题 -> 系统识别知识问答意图 -> Embedding 向量化问题 -> Qdrant Local 检索知识片段 -> 结合知识片段生成回答 -> 返回答案与知识依据。",
    )
    add_note(
        doc,
        "多轮追问流程",
        "系统读取上一轮 effective_question，继承时间范围、指标、地区和维度；当用户提出“按产品类别拆分”时追加维度，当用户提出“只看华东”时增加筛选条件。",
    )
    doc.add_heading("三、业务功能概要结构", level=2)
    add_table(
        doc,
        ["一级模块", "二级模块", "说明"],
        [
            ["登录认证", "登录、退出、登录态检查", "控制系统入口和业务接口访问。"],
            ["账户管理", "管理员列表、新增普通管理员", "初始管理员拥有新增权限，普通管理员只读。"],
            ["工作台首页", "统计概览、快捷入口、数据资产摘要", "展示系统整体运行状态。"],
            ["数据源管理", "上传数据、查看元数据、预览字段", "支撑后续 SQL 分析。"],
            ["智能分析", "自然语言提问、SQL 查询、图表、结论", "系统核心业务功能。"],
            ["多轮会话", "历史会话、消息、结果恢复、删除", "支撑连续分析体验。"],
            ["业务知识库", "知识片段新增、删除、检索、重建索引", "支撑指标口径解释和 RAG 问答。"],
            ["系统配置", "模型状态、Embedding 状态、向量库状态", "展示后端基础设施配置。"],
            ["报告导出", "HTML 报告生成", "输出分析结果留档。"],
        ],
        [1800, 3000, 4560],
    )
    doc.add_heading("四、模块定义", level=2)
    add_table(
        doc,
        ["模块", "输入", "处理", "输出"],
        [
            ["AuthRouter", "账号、密码、Cookie", "校验密码、创建/销毁会话、鉴权", "管理员信息、登录 Cookie、401/403 错误"],
            ["DatasetService", "CSV/Excel 文件", "读取文件、建表、推断字段、保存元数据", "数据集详情、字段元数据"],
            ["AnalyzerService", "自然语言问题、会话 ID、数据集 ID", "解析问题、生成 SQL、检索知识、生成结论", "AnalysisResult"],
            ["SecurityService", "SQL、授权表名", "检查只读语句、危险关键字和表范围", "安全 SQL 或异常"],
            ["KnowledgeService", "知识片段、问题文本", "关键词检索或向量检索", "相关知识片段"],
            ["VectorStoreService", "知识文本、问题文本", "Embedding、Qdrant 索引、向量查询", "向量检索结果"],
            ["Frontend API", "用户操作", "调用后端 REST API，带 Cookie", "页面状态与业务数据"],
            ["ResultChart", "chart spec 与 rows", "使用 ECharts 渲染折线图、柱状图、饼图", "可视化图表"],
        ],
        [1800, 2300, 3400, 1860],
        font_size=8.8,
    )


def add_function_structure_only(doc):
    doc.add_heading("第四部分 功能结构设计", level=1)
    doc.add_heading("一、功能结构抽取说明", level=2)
    add_para(
        doc,
        "依据需求说明书中的用户角色和用例，系统功能可以抽取为登录认证、账户管理、数据源管理、智能分析、多轮会话、业务知识库、系统配置和报告导出八类功能对象。功能结构设计以用例为起点，分离用户界面功能、后端服务功能和数据持久化对象。",
    )
    doc.add_heading("二、功能结构表", level=2)
    add_table(
        doc,
        ["功能对象", "子功能", "关联数据对象"],
        [
            ["登录认证", "登录、退出、登录态检查", "admin_users、admin_sessions"],
            ["账户管理", "管理员列表、新增普通管理员", "admin_users"],
            ["数据源管理", "上传、入库、元数据识别、预览", "datasets、dataset_columns、业务数据表"],
            ["智能分析", "问题解析、SQL 生成、安全执行、结论生成", "datasets、messages、audit_logs"],
            ["多轮会话", "创建会话、保存消息、恢复结果、删除会话", "sessions、messages"],
            ["业务知识库", "新增、删除、检索、重建向量索引", "knowledge_chunks、Qdrant collection"],
            ["图表报告", "chart spec、ECharts 渲染、HTML 导出", "messages.payload"],
        ],
        [2100, 4200, 3060],
    )


def add_er_design(doc, part="第五部分"):
    doc.add_heading(f"{part} E-R 实体设计", level=1)
    doc.add_heading("一、核心实体关系说明", level=2)
    add_para(
        doc,
        "系统实体围绕管理员、数据集、字段元数据、知识片段、会话、消息和审计日志展开。SQLite 负责保存结构化业务数据和系统元数据，Qdrant Local 负责保存知识片段向量索引。",
    )
    add_note(
        doc,
        "E-R 关系文本图",
        "admin_users 1 - N admin_sessions\nadmin_users 1 - N admin_users(created_by)\ndatasets 1 - N dataset_columns\ndatasets 1 - N knowledge_chunks\nsessions 1 - N messages\nsessions 1 - N audit_logs\nknowledge_chunks 1 - 1 Qdrant point(id 对应知识片段 id)",
    )
    doc.add_heading("二、实体定义表", level=2)
    add_table(
        doc,
        ["实体", "主键", "关键属性", "说明"],
        [
            ["admin_users", "id", "username、password_hash、is_initial_admin、created_by、created_at", "管理员账号表，区分初始管理员与普通管理员。"],
            ["admin_sessions", "token", "admin_id、created_at、expires_at", "管理员登录会话表，用于 HttpOnly Cookie 登录态。"],
            ["datasets", "id", "name、source_type、table_name、row_count、column_count", "数据集元信息。"],
            ["dataset_columns", "id", "dataset_id、name、data_type、description、sample_value、null_rate", "数据集字段元数据。"],
            ["knowledge_chunks", "id", "title、content、category、dataset_id、created_at", "业务知识片段。"],
            ["sessions", "id", "title、dataset_id、created_at、updated_at", "智能分析会话。"],
            ["messages", "id", "session_id、role、content、payload、created_at", "会话消息与分析结果。"],
            ["audit_logs", "id", "action、resource_type、resource_id、detail、status", "审计日志。"],
            ["data_demo_sales", "组合业务数据", "order_date、region、product_category、channel、sales_amount 等", "演示经营数据表。"],
        ],
        [1700, 1450, 3850, 2360],
        font_size=8.5,
    )


def add_overall_design(doc, section_title="第六部分 总体设计"):
    doc.add_heading(section_title, level=1)
    doc.add_heading("一、总体架构", level=2)
    add_note(
        doc,
        "总体架构",
        "浏览器用户 -> Vue 前端工作台 -> FastAPI 后端 API -> 智能分析服务/知识库服务/数据集服务 -> SQLite/Qdrant Local/外部模型 API。",
    )
    add_table(
        doc,
        ["层级", "组成", "职责"],
        [
            ["表现层", "Vue 3、TypeScript、ECharts", "提供登录页、工作台、数据源、智能分析、知识库、账户管理和系统配置页面。"],
            ["接口层", "FastAPI Routers", "提供 REST API，处理鉴权、参数校验、错误返回和响应模型。"],
            ["业务层", "Analyzer、Dataset、Knowledge、Auth、VectorStore 服务", "实现自然语言分析、数据导入、RAG 检索、多轮会话和管理员权限。"],
            ["数据层", "SQLite、Qdrant Local", "保存结构化数据、元数据、会话、知识片段和向量索引。"],
            ["模型层", "LLM、Embedding API", "辅助生成自然语言回答和知识语义检索。"],
        ],
        [1500, 2850, 5010],
    )
    doc.add_heading("二、后端分层设计", level=2)
    add_table(
        doc,
        ["目录/文件", "职责"],
        [
            ["main.py", "后端入口，初始化数据库，同步知识库索引，配置 CORS 和登录鉴权中间件，注册路由。"],
            ["routers/", "接口层，包括 auth、agent、datasets、knowledge、system。"],
            ["services/", "业务服务层，包括 analyzer、auth、datasets、knowledge、llm、security、vector_store。"],
            ["models.py", "Pydantic 请求与响应模型。"],
            ["db.py", "SQLite 表结构、连接管理和演示数据初始化。"],
        ],
        [2100, 7260],
    )


def add_ui_design(doc):
    doc.add_heading("第七部分 用户界面设计", level=1)
    add_table(
        doc,
        ["页面", "界面内容", "交互说明"],
        [
            ["登录页面", "账号、密码、登录按钮", "未登录访问网站时首先显示；登录成功后进入工作台。"],
            ["工作台首页", "统计卡片、快捷分析入口、数据资产摘要", "点击快捷问题可直接进入智能分析。"],
            ["智能分析页面", "左侧会话区，右侧图表和结果区", "支持输入问题、多轮追问、查看 SQL、导出报告。"],
            ["数据源页面", "上传区域、数据集列表、字段表", "支持 CSV/Excel 上传和数据预览。"],
            ["业务知识库页面", "新增知识片段表单、知识片段列表", "支持新增和删除知识，自动同步索引。"],
            ["账户管理页面", "管理员列表、新增管理员表单", "初始管理员可新增普通管理员，普通管理员显示无权限提示。"],
            ["系统配置页面", "LLM、Embedding、向量库状态", "只展示配置状态，不展示 API Key 明文。"],
        ],
        [1700, 4000, 3660],
    )
    doc.add_heading("二、界面风格要求", level=2)
    for item in [
        "整体采用后台管理系统风格，左侧导航、顶部数据集选择器、主内容区展示业务功能。",
        "主色调使用深蓝、科技蓝和青绿色，突出数据智能体和企业分析场景。",
        "图表区域使用 ECharts，支持多系列图例滚动和横向缩放。",
        "登录页面必须简洁，不展示初始管理员账号提示，避免暴露账号信息。",
    ]:
        add_bullet(doc, item)


def add_algorithm_design(doc):
    doc.add_heading("第八部分 算法设计", level=1)
    doc.add_heading("一、运行环境说明", level=2)
    add_table(
        doc,
        ["运行对象", "环境要求"],
        [
            ["前端", "Node.js、Vite、Vue 3、ECharts，默认端口 5173。"],
            ["后端", "Python、FastAPI、Uvicorn、Pydantic、httpx，默认端口 8000。"],
            ["结构化存储", "SQLite 本地文件 backend/storage/data_agent.db。"],
            ["向量存储", "Qdrant Local，本地目录 backend/storage/qdrant。"],
            ["模型服务", "OpenAI 兼容 LLM 接口和 Embedding 接口，由 backend/.env 配置。"],
        ],
        [1800, 7560],
    )
    doc.add_heading("二、数据集制作说明", level=2)
    add_table(
        doc,
        ["数据来源", "字段", "用途"],
        [
            ["演示经营数据", "order_date、region、product_category、channel、sales_amount、order_count、profit、complaint_count、visits、conversions", "用于演示销售额趋势、投诉率、转化率、利润等经营分析。"],
            ["用户上传 CSV/Excel", "系统自动读取首行字段并推断类型", "用于扩展自定义数据集分析。"],
            ["业务知识片段", "title、content、category、dataset_id", "用于指标口径、字段解释和业务规则问答。"],
        ],
        [1900, 4300, 3160],
        font_size=8.8,
    )
    doc.add_heading("三、核心模块要求", level=2)
    add_table(
        doc,
        ["模块", "算法/规则", "输出"],
        [
            ["登录认证", "PBKDF2-SHA256 加盐哈希校验密码；随机 token 写入 admin_sessions；HttpOnly Cookie 保存 token。", "管理员登录态"],
            ["指标识别", "根据关键词识别销售额、订单数、利润、投诉率、转化率等指标。", "metric_sql、metric_label"],
            ["时间识别", "正则识别近 N 天、近 N 周、近 N 月、本月、本年。", "SQL WHERE 时间过滤条件"],
            ["维度识别", "根据地区、区域、产品、展品、品类、渠道等关键词识别拆分维度。", "x_field、series_fields"],
            ["多轮合并", "读取上一轮 effective_question，追问时继承时间、指标和维度；显式只看/去掉时替换条件。", "新的 effective_question"],
            ["SQL 安全", "校验只读语句、危险关键字和授权表范围。", "safe_sql 或异常"],
            ["RAG 检索", "Embedding 问题文本，Qdrant 查询相似知识片段，并加入关键词加权。", "knowledge_refs"],
            ["图表生成", "根据时间趋势和维度关系选择 line/bar/pie，并生成 chart spec。", "前端 ECharts 配置数据"],
        ],
        [1600, 5850, 1910],
        font_size=8.5,
    )


def add_deployment(doc):
    doc.add_heading("第九部分 运行环境和部署", level=1)
    add_table(
        doc,
        ["节点", "部署内容", "说明"],
        [
            ["浏览器客户端", "访问 http://127.0.0.1:5173", "用户登录和使用系统的入口。"],
            ["前端服务", "Vite Dev Server / 静态资源", "开发环境默认端口 5173，通过 /api 代理后端。"],
            ["后端服务", "FastAPI + Uvicorn", "默认端口 8000，提供 /api/v1 接口。"],
            ["SQLite", "本地数据库文件", "保存数据集、元数据、会话、知识片段、管理员和审计日志。"],
            ["Qdrant Local", "本地向量索引目录", "保存知识库向量。"],
            ["外部模型服务", "LLM 与 Embedding API", "由后端 .env 配置访问。"],
        ],
        [1700, 3000, 4660],
    )
    add_note(
        doc,
        "部署命令",
        "后端：cd backend && .venv\\Scripts\\python.exe -m uvicorn app.main:app --host 127.0.0.1 --port 8000\n前端：cd frontend && npm run dev\n一键启动：scripts\\start-dev.cmd 或 .\\scripts\\start-dev.ps1",
    )


def add_doc1_content(doc):
    add_common_intro(doc)
    add_project_overview(doc)
    add_constraints(doc)
    add_function_structure_only(doc)
    add_er_design(doc, "第五部分")


def add_doc2_content(doc):
    add_common_intro(doc)
    add_project_overview(doc)
    add_constraints(doc)
    add_solution_design(doc)
    add_er_design(doc, "第五部分")
    add_overall_design(doc, "第六部分 总体设计")
    add_ui_design(doc)
    add_algorithm_design(doc)
    add_deployment(doc)


def add_contents(doc, full: bool):
    doc.add_heading("目 录", level=1)
    entries = [
        "第一部分 引言",
        "第二部分 项目概述",
        "  一、项目目标描述",
        "  二、项目功能描述",
        "第三部分 设计约束",
    ]
    if full:
        entries.extend(
            [
                "第四部分 数据智能体服务系统方案设计",
                "  一、安全设计",
                "  二、数据智能体服务系统相关业务流程",
                "  三、业务功能概要结构",
                "  四、模块定义",
                "第五部分 E-R 实体设计",
                "第六部分 总体设计",
                "第七部分 用户界面设计",
                "第八部分 算法设计",
                "  一、运行环境说明",
                "  二、数据集制作说明",
                "  三、核心模块要求",
                "第九部分 运行环境和部署",
            ]
        )
    else:
        entries.extend(
            [
                "第四部分 功能结构设计",
                "第五部分 E-R 实体设计",
            ]
        )
    for item in entries:
        add_para(doc, item, size=10.5, after=2)
    doc.add_page_break()


def build(path: Path, subtitle: str, full: bool):
    doc = Document()
    configure_doc(doc, subtitle)
    add_cover(doc, subtitle)
    add_contents(doc, full)
    if full:
        add_doc2_content(doc)
    else:
        add_doc1_content(doc)
    doc.core_properties.title = f"数据智能体服务系统{subtitle}"
    doc.core_properties.subject = "概要设计"
    doc.core_properties.author = "数据智能体项目组"
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


if __name__ == "__main__":
    print(build(OUT1, "概要设计说明书（1/2）", full=False))
    print(build(OUT2, "概要设计说明书（2/2 完整版）", full=True))
