from __future__ import annotations

from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "assets" / "overview_design"
OUT = DOCS / "《数据智能体服务系统》概要设计-2_图文增强版.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
GRAY = RGBColor(90, 90, 90)

FONT_PATHS = [
    Path("C:/Windows/Fonts/msyh.ttc"),
    Path("C:/Windows/Fonts/simhei.ttf"),
    Path("C:/Windows/Fonts/simsun.ttc"),
]


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [Path("C:/Windows/Fonts/msyhbd.ttc")] if bold else []
    candidates += FONT_PATHS
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


F_TITLE = font(42, True)
F_H1 = font(34, True)
F_H2 = font(28, True)
F_BODY = font(24)
F_SMALL = font(20)
F_TINY = font(18)


def text_size(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.ImageFont) -> tuple[int, int]:
    box = draw.textbbox((0, 0), text, font=fnt)
    return box[2] - box[0], box[3] - box[1]


def wrap_text(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.ImageFont, max_width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for char in text:
        test = current + char
        if draw.textlength(test, font=fnt) <= max_width or not current:
            current = test
        else:
            lines.append(current)
            current = char
    if current:
        lines.append(current)
    return lines


def draw_text_center(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    fnt: ImageFont.ImageFont,
    fill=(24, 45, 78),
    line_gap: int = 6,
) -> None:
    x1, y1, x2, y2 = box
    lines: list[str] = []
    for part in text.split("\n"):
        lines.extend(wrap_text(draw, part, fnt, x2 - x1 - 28))
    heights = [text_size(draw, line, fnt)[1] for line in lines]
    total_h = sum(heights) + line_gap * max(0, len(lines) - 1)
    y = y1 + (y2 - y1 - total_h) // 2
    for line, h in zip(lines, heights):
        w, _ = text_size(draw, line, fnt)
        draw.text((x1 + (x2 - x1 - w) // 2, y), line, font=fnt, fill=fill)
        y += h + line_gap


def draw_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    fill=(255, 255, 255),
    outline=(88, 143, 184),
    text_fill=(24, 45, 78),
    radius: int = 22,
    width: int = 3,
    fnt: ImageFont.ImageFont = F_BODY,
) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)
    draw_text_center(draw, box, text, fnt, text_fill)


def arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    fill=(66, 99, 139),
    width: int = 4,
    dashed: bool = False,
) -> None:
    x1, y1 = start
    x2, y2 = end
    if dashed:
        steps = 18
        for i in range(steps):
            if i % 2 == 0:
                sx = x1 + (x2 - x1) * i / steps
                sy = y1 + (y2 - y1) * i / steps
                ex = x1 + (x2 - x1) * (i + 1) / steps
                ey = y1 + (y2 - y1) * (i + 1) / steps
                draw.line((sx, sy, ex, ey), fill=fill, width=width)
    else:
        draw.line((x1, y1, x2, y2), fill=fill, width=width)
    # Arrow head
    import math

    angle = math.atan2(y2 - y1, x2 - x1)
    size = 18
    p1 = (x2, y2)
    p2 = (x2 - size * math.cos(angle - 0.42), y2 - size * math.sin(angle - 0.42))
    p3 = (x2 - size * math.cos(angle + 0.42), y2 - size * math.sin(angle + 0.42))
    draw.polygon([p1, p2, p3], fill=fill)


def add_caption_text(draw: ImageDraw.ImageDraw, text: str, y: int, w: int) -> None:
    draw.text((70, y), text, font=F_SMALL, fill=(80, 96, 119))


def canvas(w: int, h: int, title: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (w, h), (248, 251, 253))
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((30, 30, w - 30, h - 30), radius=30, fill=(255, 255, 255), outline=(213, 224, 235), width=3)
    draw.text((70, 58), title, font=F_TITLE, fill=(11, 37, 69))
    draw.line((70, 120, w - 70, 120), fill=(221, 232, 242), width=4)
    return img, draw


def make_architecture(path: Path) -> None:
    img, draw = canvas(2400, 1540, "数据智能体服务系统总体架构图")
    layers = [
        ("用户与展示层", 170, [(230, "浏览器客户端\nVue3 + TypeScript"), (600, "智能分析页面\n对话/图表/表格"), (970, "数据源页面\nCSV上传/字段预览"), (1340, "知识库页面\n知识片段管理"), (1710, "账户与配置页面\n登录/管理员/模型状态")]),
        ("API 接入层", 450, [(320, "FastAPI 应用\nREST API"), (740, "登录鉴权中间件\nCookie Session"), (1160, "路由层\nagent/datasets/knowledge/auth/system"), (1580, "响应封装\nJSON/HTML报告")]),
        ("智能体与业务服务层", 730, [(210, "Analyzer\n意图识别/SQL规划"), (560, "Dataset Service\n数据集导入/元数据"), (910, "Knowledge Service\n片段管理/RAG检索"), (1260, "LLM Adapter\n结论润色/知识问答"), (1610, "Auth Service\n密码哈希/权限校验")]),
        ("数据与模型基础设施层", 1010, [(260, "SQLite\n业务数据/元数据/会话"), (730, "Qdrant Local\n知识向量索引"), (1200, "Embedding API\n文本向量化"), (1670, "LLM API\nOpenAI兼容接口")]),
    ]
    colors = [(234, 246, 252), (241, 247, 255), (246, 250, 244), (255, 248, 236)]
    for idx, (name, y, boxes) in enumerate(layers):
        draw.rounded_rectangle((90, y, 2310, y + 210), radius=26, fill=colors[idx], outline=(198, 215, 230), width=2)
        draw.text((125, y + 22), name, font=F_H2, fill=(31, 77, 120))
        for x, label in boxes:
            draw_box(draw, (x, y + 68, x + 300, y + 178), label, fill=(255, 255, 255), outline=(72, 145, 177), fnt=F_SMALL)
    for y in [380, 660, 940]:
        arrow(draw, (1200, y), (1200, y + 65), fill=(73, 122, 163), width=5)
    # Bottom-layer components are intentionally not connected by horizontal arrows
    # to avoid crossing labels after the diagram is scaled into Word.
    add_caption_text(draw, "说明：前端仅负责交互展示，API Key、SQL执行、向量检索和模型调用全部在后端完成。", 1420, 2400)
    img.save(path)


def make_function_modules(path: Path) -> None:
    img, draw = canvas(2300, 1420, "业务功能模块结构图")
    root = (790, 165, 1510, 255)
    draw_box(draw, root, "数据智能体服务系统", fill=(25, 91, 140), outline=(25, 91, 140), text_fill=(255, 255, 255), fnt=F_H2)
    modules = [
        ("工作台", ["运行概览", "数据集统计", "知识库状态", "会话摘要"]),
        ("智能分析", ["自然语言提问", "多轮上下文", "SQL生成与校验", "图表/洞察/报告"]),
        ("数据源管理", ["CSV上传", "字段识别", "数据预览", "数据集切换"]),
        ("业务知识库", ["知识片段新增", "向量索引同步", "知识问答", "片段删除"]),
        ("历史会话", ["会话列表", "旧对话恢复", "分析结果回放", "会话删除"]),
        ("账户管理", ["登录退出", "初始管理员", "新增管理员", "权限限制"]),
        ("系统配置", ["LLM状态", "Embedding状态", "SQLite/Qdrant状态", "API Key后端配置"]),
    ]
    positions = [(220, 380), (530, 380), (840, 380), (1150, 380), (1460, 380), (1770, 380), (2080, 380)]
    for (title, items), (x, y) in zip(modules, positions):
        box = (x - 105, y, x + 185, y + 115)
        arrow(draw, ((root[0] + root[2]) // 2, root[3]), (x + 40, y), fill=(90, 126, 166), width=4)
        draw_box(draw, box, title, fill=(236, 247, 252), outline=(58, 151, 190), fnt=F_BODY)
        iy = y + 150
        for item in items:
            draw.rounded_rectangle((x - 95, iy, x + 175, iy + 56), radius=15, fill=(255, 255, 255), outline=(210, 224, 235), width=2)
            draw_text_center(draw, (x - 95, iy, x + 175, iy + 56), item, F_TINY, fill=(33, 55, 84))
            iy += 72
    add_caption_text(draw, "说明：功能结构来自需求用例与当前实现模块，智能分析与知识库共同构成数据智能体核心能力。", 1320, 2300)
    img.save(path)


def make_sequence(path: Path, title: str, actors: list[str], messages: list[tuple[int, int, str, bool]]) -> None:
    w, h = 2500, 1600
    img, draw = canvas(w, h, title)
    left, top = 175, 175
    gap = (w - 2 * left) // (len(actors) - 1)
    xs = [left + i * gap for i in range(len(actors))]
    for x, actor in zip(xs, actors):
        draw_box(draw, (x - 110, top, x + 110, top + 76), actor, fill=(236, 247, 252), outline=(58, 151, 190), fnt=F_SMALL)
        # lifeline
        y = top + 88
        while y < h - 90:
            draw.line((x, y, x, min(y + 18, h - 90)), fill=(168, 184, 202), width=3)
            y += 34
    y = top + 145
    for src, dst, label, ret in messages:
        x1, x2 = xs[src], xs[dst]
        if src == dst:
            loop_x = min(x1 + 150, w - 150)
            draw.line((x1, y, loop_x, y, loop_x, y + 42, x1, y + 42), fill=(69, 104, 142), width=4)
            arrow(draw, (x1 + 35, y + 42), (x1, y + 42), fill=(69, 104, 142), width=4)
            label_box = (x1 + 28, y - 50, loop_x + 110, y - 8)
        else:
            arrow(draw, (x1, y), (x2, y), fill=(69, 104, 142), width=4, dashed=ret)
            label_box = (min(x1, x2) + 16, y - 48, max(x1, x2) - 16, y - 8)
        draw.rounded_rectangle(label_box, radius=10, fill=(255, 255, 255), outline=(230, 236, 242), width=1)
        draw_text_center(draw, label_box, label, F_TINY, fill=(31, 58, 91))
        y += 105
    draw.text((90, h - 72), "注：实线表示请求/调用，虚线表示返回。", font=F_SMALL, fill=(80, 96, 119))
    img.save(path)


def make_er(path: Path) -> None:
    img, draw = canvas(2500, 1780, "E-R实体关系图")
    entities = {
        "admin_users": ((120, 210, 600, 440), ["id PK", "username UNIQUE", "password_hash", "is_initial_admin", "created_by FK"]),
        "admin_sessions": ((800, 210, 1280, 390), ["token PK", "admin_id FK", "created_at", "expires_at"]),
        "audit_logs": ((1500, 210, 1980, 440), ["id PK", "action", "resource_type", "resource_id", "status"]),
        "datasets": ((120, 630, 600, 900), ["id PK", "name", "source_type", "table_name UNIQUE", "row_count", "column_count"]),
        "dataset_columns": ((800, 630, 1280, 880), ["id PK", "dataset_id FK", "name", "data_type", "description", "sample_value"]),
        "data_*": ((1500, 630, 1980, 880), ["order_date", "region", "product_category", "channel", "sales/profit/orders..."]),
        "knowledge_chunks": ((120, 1100, 600, 1350), ["id PK", "title", "content", "category", "dataset_id FK", "created_at"]),
        "sessions": ((800, 1100, 1280, 1320), ["id PK", "title", "dataset_id FK", "created_at", "updated_at"]),
        "messages": ((1500, 1100, 1980, 1350), ["id PK", "session_id FK", "role", "content", "payload", "created_at"]),
    }
    for name, (box, fields) in entities.items():
        draw.rounded_rectangle(box, radius=22, fill=(255, 255, 255), outline=(71, 134, 176), width=3)
        x1, y1, x2, y2 = box
        draw.rounded_rectangle((x1, y1, x2, y1 + 52), radius=22, fill=(25, 91, 140), outline=(25, 91, 140), width=0)
        draw_text_center(draw, (x1, y1, x2, y1 + 52), name, F_SMALL, fill=(255, 255, 255))
        fy = y1 + 70
        for field in fields:
            draw.text((x1 + 22, fy), field, font=F_TINY, fill=(34, 56, 88))
            fy += 31

    def connect(a: str, b: str, label: str, start_side="right", end_side="left"):
        box_a = entities[a][0]
        box_b = entities[b][0]
        if start_side == "right":
            s = (box_a[2], (box_a[1] + box_a[3]) // 2)
        elif start_side == "bottom":
            s = ((box_a[0] + box_a[2]) // 2, box_a[3])
        else:
            s = (box_a[0], (box_a[1] + box_a[3]) // 2)
        if end_side == "left":
            e = (box_b[0], (box_b[1] + box_b[3]) // 2)
        elif end_side == "top":
            e = ((box_b[0] + box_b[2]) // 2, box_b[1])
        else:
            e = (box_b[2], (box_b[1] + box_b[3]) // 2)
        arrow(draw, s, e, fill=(96, 125, 160), width=4)
        mx, my = (s[0] + e[0]) // 2, (s[1] + e[1]) // 2
        draw.rounded_rectangle((mx - 85, my - 23, mx + 85, my + 23), radius=12, fill=(248, 251, 253), outline=(220, 230, 240))
        draw_text_center(draw, (mx - 85, my - 23, mx + 85, my + 23), label, F_TINY, fill=(64, 82, 110))

    connect("admin_users", "admin_sessions", "1:N")
    connect("datasets", "dataset_columns", "1:N")
    connect("dataset_columns", "data_*", "字段映射")
    connect("datasets", "knowledge_chunks", "1:N", "bottom", "top")
    connect("datasets", "sessions", "1:N", "right", "left")
    connect("sessions", "messages", "1:N")
    draw.rounded_rectangle((1375, 375, 1615, 425), radius=14, fill=(248, 251, 253), outline=(220, 230, 240))
    draw_text_center(draw, (1375, 375, 1615, 425), "记录系统操作", F_TINY, fill=(64, 82, 110))
    draw.text((90, 1695), "说明：业务数据表 data_* 由上传CSV动态创建，datasets.table_name 保存对应表名；messages.payload 保存图表、SQL、结论等分析结果。", font=F_SMALL, fill=(80, 96, 119))
    img.save(path)


def make_algorithm(path: Path) -> None:
    img, draw = canvas(2300, 980, "智能分析与RAG核心算法流程图")
    steps = [
        ("用户问题", "自然语言\n+ 数据集"),
        ("上下文合并", "继承历史会话\n时间/指标/维度"),
        ("意图识别", "数据分析\n或知识问答"),
        ("SQL/RAG规划", "只读SQL\n向量检索"),
        ("安全执行", "SQL校验\nSQLite查询"),
        ("模型增强", "LLM润色\n知识依据"),
        ("结果生成", "图表/表格\n洞察/报告"),
        ("历史保存", "messages\npayload"),
    ]
    x = 90
    y = 280
    bw, bh, gap = 230, 150, 45
    for idx, (title, body) in enumerate(steps):
        fill = (236, 247, 252) if idx % 2 == 0 else (246, 250, 244)
        draw_box(draw, (x, y, x + bw, y + bh), f"{title}\n{body}", fill=fill, outline=(72, 145, 177), fnt=F_SMALL)
        if idx < len(steps) - 1:
            arrow(draw, (x + bw, y + bh // 2), (x + bw + gap - 6, y + bh // 2), fill=(73, 122, 163), width=5)
        x += bw + gap
    draw.rounded_rectangle((120, 590, 2180, 820), radius=24, fill=(255, 248, 236), outline=(231, 211, 166), width=3)
    notes = [
        "安全策略：仅允许SELECT类只读查询，禁止DROP/DELETE/UPDATE/INSERT等危险关键字。",
        "多轮策略：读取上一轮effective_question，除非用户显式替换，否则继承时间范围、指标和维度。",
        "降级策略：LLM或向量服务不可用时，仍可使用规则SQL与本地知识关键词匹配返回基础结果。",
    ]
    yy = 625
    for note in notes:
        draw.text((160, yy), f"• {note}", font=F_SMALL, fill=(57, 73, 96))
        yy += 58
    img.save(path)


def make_ui_flow(path: Path) -> None:
    img, draw = canvas(2300, 1120, "用户界面导航结构图")
    boxes = [
        ("登录页", 100, 260, "账号密码登录\nHttpOnly Cookie"),
        ("工作台", 520, 170, "运行概览\n数据/知识/会话统计"),
        ("智能分析", 940, 120, "多轮问答\n图表/表格/SQL/报告"),
        ("数据源", 940, 380, "CSV上传\n字段预览/数据集切换"),
        ("知识库", 1360, 120, "知识新增\n向量同步/删除"),
        ("账户管理", 1360, 380, "初始管理员\n新增管理员/权限提示"),
        ("系统配置", 1780, 250, "LLM/Embedding\nSQLite/Qdrant状态"),
    ]
    for title, x, y, desc in boxes:
        draw_box(draw, (x, y, x + 300, y + 135), f"{title}\n{desc}", fill=(255, 255, 255), outline=(72, 145, 177), fnt=F_SMALL)
    arrow(draw, (400, 327), (520, 245), width=5)
    for start, end in [((820, 237), (940, 187)), ((820, 237), (940, 447)), ((1240, 187), (1360, 187)), ((1240, 447), (1360, 447)), ((1660, 447), (1780, 317)), ((1660, 187), (1780, 317))]:
        arrow(draw, start, end, width=4)
    draw.rounded_rectangle((100, 650, 2200, 900), radius=24, fill=(239, 249, 251), outline=(202, 224, 234), width=3)
    draw.text((145, 695), "界面设计原则", font=F_H2, fill=(31, 77, 120))
    principles = [
        "先登录后进入系统，避免未授权访问业务数据和系统配置。",
        "智能分析页保持“左侧对话过程 + 右侧报告结果”的布局，适合答辩演示分析路径。",
        "图表组件支持多系列维度，二轮追问按展品/地区拆分时保留上一轮地区维度。",
        "知识库与历史会话提供删除能力，便于测试数据和演示数据清理。",
    ]
    yy = 750
    for item in principles:
        draw.text((175, yy), f"• {item}", font=F_SMALL, fill=(57, 73, 96))
        yy += 44
    img.save(path)


def make_diagrams() -> dict[str, Path]:
    ASSETS.mkdir(parents=True, exist_ok=True)
    paths = {
        "architecture": ASSETS / "01_system_architecture.png",
        "modules": ASSETS / "02_function_modules.png",
        "sequence_analysis": ASSETS / "03_sequence_analysis.png",
        "sequence_knowledge": ASSETS / "04_sequence_knowledge.png",
        "er": ASSETS / "05_er_model.png",
        "algorithm": ASSETS / "06_algorithm_pipeline.png",
        "ui": ASSETS / "07_ui_flow.png",
    }
    make_architecture(paths["architecture"])
    make_function_modules(paths["modules"])
    make_sequence(
        paths["sequence_analysis"],
        "核心功能时序图：智能分析与图表生成",
        ["管理员", "前端Vue", "API路由", "会话服务", "分析Agent", "SQLite", "知识检索", "LLM"],
        [
            (0, 1, "输入问题并选择数据集", False),
            (1, 2, "POST /api/v1/agent/chat", False),
            (2, 3, "读取/创建会话", False),
            (3, 4, "加载历史上下文", False),
            (4, 5, "读取数据集元数据", False),
            (4, 4, "识别指标/时间/维度并生成SQL", False),
            (4, 5, "校验只读SQL并查询聚合结果", False),
            (4, 6, "检索业务知识片段", False),
            (4, 7, "调用LLM润色洞察", False),
            (4, 3, "保存会话消息与结果payload", False),
            (2, 1, "返回图表、表格、SQL、洞察", True),
        ],
    )
    make_sequence(
        paths["sequence_knowledge"],
        "核心功能时序图：知识库问答与RAG检索",
        ["管理员", "前端Vue", "API路由", "Agent服务", "Embedding", "Qdrant", "LLM", "SQLite"],
        [
            (0, 1, "提出业务规则/指标口径问题", False),
            (1, 2, "POST /api/v1/agent/chat", False),
            (2, 3, "判断为knowledge_qa", False),
            (3, 7, "读取历史对话上下文", False),
            (3, 4, "问题文本向量化", False),
            (3, 5, "向量相似度检索知识片段", False),
            (3, 6, "拼接知识上下文生成回答", False),
            (3, 7, "保存回答与引用片段", False),
            (2, 1, "返回回答、引用依据、执行模式", True),
        ],
    )
    make_er(paths["er"])
    make_algorithm(paths["algorithm"])
    make_ui_flow(paths["ui"])
    return paths


def set_run_font(run, size=None, bold=None, color=None, name="Calibri", east_asia="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def pformat(p, before=0, after=6, line=1.10, align=None):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align


def configure_doc(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10
    header = section.header.paragraphs[0]
    header.text = "数据智能体服务系统 | 概要设计说明书"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, 9, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.text = "概要设计-2 图文增强版"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run_font(run, 9, color=GRAY)


def add_para(doc, text="", size=11, bold=False, color=None, style=None, before=0, after=6, align=None):
    p = doc.add_paragraph(style=style)
    pformat(p, before=before, after=after, align=align)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    pformat(p, after=4, line=1.167)
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    pformat(p, after=0)
    r = p.add_run(str(text))
    set_run_font(r, size=9.5, bold=bold, color=NAVY if bold else None)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill:
        shade_cell(cell, fill)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, fill="F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return table


def add_figure(doc, image_path: Path, caption: str, width: float = 6.45):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    run = cap.add_run(caption)
    set_run_font(run, size=9.5, color=GRAY, bold=True)


def add_cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    add_para(doc, "数据智能体服务系统", size=28, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, "概要设计说明书（2/2 完整版·图文增强版）", size=18, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=30)
    add_para(doc, "Software Architecture & High-Level Design", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=50)
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["文档阶段", "概要设计-2：UI设计、技术架构设计与总体设计"],
            ["适用范围", "数据智能体服务系统的开发、答辩、部署与后续详细设计"],
            ["主要图件", "系统架构图、功能模块图、核心功能时序图、E-R实体图、算法流程图、界面导航图"],
            ["版本", "V1.1 图文增强版"],
        ],
        [1.8, 4.5],
    )
    doc.add_page_break()


def add_contents(doc):
    doc.add_heading("目 录", level=1)
    entries = [
        "第一部分 引言",
        "第二部分 项目概述",
        "  一、项目目标描述",
        "  二、项目功能描述",
        "第三部分 设计约束",
        "第四部分 数据智能体服务系统方案设计",
        "  一、安全设计",
        "  二、数据智能体服务系统相关业务流程",
        "  三、业务功能概要结构",
        "  四、模块定义",
        "第五部分 E-R 实体设计",
        "第六部分 总体设计",
        "第七部分 用户界面设计",
        "第八部分 算法设计",
        "  一、运行环境说明",
        "  二、数据集制作说明",
        "  三、核心模块要求",
        "第九部分 运行环境和部署",
    ]
    for entry in entries:
        add_para(doc, entry, size=10.5, after=2)
    doc.add_page_break()


def build_doc(paths: dict[str, Path]) -> Path:
    doc = Document()
    configure_doc(doc)
    add_cover(doc)
    add_contents(doc)

    doc.add_heading("第一部分 引言", level=1)
    doc.add_heading("一、说明", level=2)
    add_para(doc, "本文档在项目需求说明书与概要设计实训模板基础上编写，用于描述数据智能体服务系统的总体方案、功能结构、核心流程、实体关系、界面结构、算法策略与部署环境。本文档作为后续详细设计、编码实现、测试验收与答辩说明的依据。")
    doc.add_heading("二、定义", level=2)
    add_table(
        doc,
        ["术语", "说明"],
        [
            ["数据智能体", "面向业务数据的智能问答与分析服务，能够理解自然语言问题、生成查询计划、执行查询并返回图表和洞察。"],
            ["RAG", "检索增强生成，通过向量检索业务知识片段，并将检索结果作为大模型回答依据。"],
            ["多轮会话", "系统保存用户与助手历史消息，在追问时继承上一轮问题中的时间、指标、维度等上下文。"],
            ["只读SQL", "仅允许查询类SQL，不允许修改、删除或破坏数据库结构。"],
        ],
        [1.4, 4.9],
    )
    doc.add_heading("三、参考资料", level=2)
    for item in [
        "《数据智能体服务系统》项目需求说明书",
        "D03_概要设计-1.pdf：功能结构与实体结构设计要求",
        "D03_概要设计-2.pdf：时序图、UI设计、技术架构与总体设计要求",
        "项目源码：frontend、backend、docs、scripts 目录",
    ]:
        add_bullet(doc, item)

    doc.add_heading("第二部分 项目概述", level=1)
    doc.add_heading("一、项目目标描述", level=2)
    add_para(doc, "本项目建设一个轻量级数据智能体服务系统，面向企业经营演示数据和用户上传CSV数据，提供自然语言数据分析、知识库问答、多轮会话、历史回看、图表生成、报告导出和管理员登录管理能力。系统目标是在本地即可完成数据接入、业务知识维护和智能分析演示，同时保留可替换的大模型与Embedding接口。")
    doc.add_heading("二、项目功能描述", level=2)
    add_table(
        doc,
        ["功能域", "功能描述", "设计重点"],
        [
            ["登录与权限", "使用账号密码登录，初始管理员可新增普通管理员，普通管理员新增管理员时提示无权限。", "HttpOnly Cookie、密码哈希、权限校验"],
            ["数据源管理", "支持CSV上传、字段元数据解析、演示数据初始化、数据集详情查看。", "结构化存储、字段描述、动态数据表"],
            ["智能分析", "用户用自然语言提出分析问题，系统识别指标、维度、时间范围，生成安全SQL并返回图表和结论。", "意图识别、只读SQL、安全执行、多系列图表"],
            ["多轮会话", "保存历史会话与消息，支持打开旧对话、继续追问和删除历史。", "上下文继承、payload回放、会话生命周期"],
            ["知识库问答", "维护业务口径、指标解释、业务规则等知识片段，结合向量检索实现知识问答。", "Qdrant Local、Embedding、RAG引用依据"],
            ["系统配置", "展示LLM、Embedding、SQLite、Qdrant等基础设施状态，API Key仅后端可见。", "配置安全、状态检测、降级策略"],
        ],
        [1.15, 3.25, 1.9],
    )

    doc.add_heading("第三部分 设计约束", level=1)
    add_table(
        doc,
        ["约束类型", "约束内容", "设计影响"],
        [
            ["技术栈", "前端采用Vue 3 + TypeScript + Vite，后端采用FastAPI，数据库采用SQLite，向量库采用Qdrant Local。", "保证轻量部署，适合三周实训项目和本地演示。"],
            ["安全性", "API Key不返回浏览器；数据库查询必须经过只读SQL校验；登录态使用HttpOnly Cookie。", "敏感信息留在后端，降低误操作风险。"],
            ["模型依赖", "LLM与Embedding均使用OpenAI兼容接口，允许配置DeepSeek、硅基流动等服务。", "模型可替换，未配置时提供规则化降级。"],
            ["数据规模", "面向课程项目与企业演示数据，优先支持中小规模CSV与本地SQLite查询。", "以可靠演示和结构清晰为目标，不设计分布式数仓。"],
            ["可维护性", "前后端分层、路由按业务划分、服务层封装算法与外部调用。", "方便在答辩中解释，也方便后续扩展。"],
        ],
        [1.15, 3.25, 1.9],
    )

    doc.add_heading("第四部分 数据智能体服务系统方案设计", level=1)
    doc.add_heading("一、安全设计", level=2)
    add_para(doc, "系统安全设计围绕身份认证、接口访问、敏感配置、SQL执行和管理员权限五个方面展开。登录后才能访问系统核心接口，API Key仅由后端读取，普通管理员不具备新增管理员权限。")
    add_table(
        doc,
        ["安全点", "设计方案", "说明"],
        [
            ["身份认证", "账号密码登录，后端生成随机Session Token并写入HttpOnly Cookie。", "浏览器脚本无法读取Token，降低泄露风险。"],
            ["密码保存", "PBKDF2-SHA256加盐哈希保存。", "数据库不保存明文密码。"],
            ["权限控制", "初始管理员is_initial_admin=1，仅初始管理员可新增管理员。", "满足“只有初始管理员可以加入其他管理员”的需求。"],
            ["SQL安全", "分析SQL必须通过只读校验，禁止危险关键字。", "避免自然语言生成破坏性SQL。"],
            ["配置安全", "LLM_API_KEY、EMBEDDING_API_KEY等保存在backend/.env。", "前端只展示配置状态，不展示密钥。"],
        ],
        [1.25, 3.25, 1.8],
    )
    doc.add_heading("二、数据智能体服务系统相关业务流程", level=2)
    add_para(doc, "根据概要设计-2要求，核心功能采用时序图表达。系统核心流程包括智能分析与图表生成、知识库问答与RAG检索。")
    add_figure(doc, paths["sequence_analysis"], "图4-1 核心功能时序图：智能分析与图表生成")
    add_figure(doc, paths["sequence_knowledge"], "图4-2 核心功能时序图：知识库问答与RAG检索")
    doc.add_heading("三、业务功能概要结构", level=2)
    add_para(doc, "系统功能从用户入口、业务数据、智能分析、知识检索和系统管理五条主线展开。功能模块图如下。")
    add_figure(doc, paths["modules"], "图4-3 业务功能模块结构图")
    doc.add_heading("四、模块定义", level=2)
    add_table(
        doc,
        ["模块", "主要职责", "关键接口/文件"],
        [
            ["认证模块", "登录、退出、获取当前管理员、新增管理员、权限校验。", "routers/auth.py、services/auth.py"],
            ["数据集模块", "CSV上传、演示数据管理、字段元数据维护、数据预览。", "routers/datasets.py、services/datasets.py"],
            ["智能分析模块", "自然语言解析、查询计划生成、SQL执行、图表规格生成、结论输出。", "routers/agent.py、services/analyzer.py"],
            ["知识库模块", "知识片段新增/删除、向量索引同步、知识检索。", "routers/knowledge.py、services/vector_store.py"],
            ["会话模块", "会话创建、消息保存、历史会话读取、删除历史对话。", "sessions、messages表"],
            ["系统配置模块", "展示数据库、向量库、LLM、Embedding配置状态。", "routers/system.py、config.py"],
            ["前端展示模块", "登录页、工作台、智能分析、数据源、知识库、账户管理、系统配置。", "frontend/src/App.vue、ResultChart.vue"],
        ],
        [1.25, 3.2, 1.85],
    )

    doc.add_heading("第五部分 E-R 实体设计", level=1)
    add_para(doc, "实体结构设计从数据集、字段元数据、知识片段、会话消息、管理员与审计日志出发。业务数据表由上传CSV动态生成，数据集表保存其真实表名。")
    add_figure(doc, paths["er"], "图5-1 E-R实体关系图")
    add_table(
        doc,
        ["实体", "含义", "关键关系"],
        [
            ["datasets", "数据集主表，记录数据集名称、来源、真实表名、行列数。", "与dataset_columns、knowledge_chunks、sessions是一对多关系。"],
            ["dataset_columns", "字段元数据表，保存字段名、类型、描述、样例值。", "从属于datasets。"],
            ["knowledge_chunks", "业务知识片段表，保存指标口径、业务规则、数据字典。", "可绑定数据集，也可作为全局知识。"],
            ["sessions/messages", "会话和消息表，保存多轮对话、助手回答和payload结果。", "sessions与messages是一对多关系。"],
            ["admin_users/admin_sessions", "管理员账号与登录会话表。", "管理员与会话Token是一对多关系。"],
            ["audit_logs", "审计日志表，记录关键动作、资源类型和执行状态。", "用于后续追踪系统操作。"],
        ],
        [1.35, 2.75, 2.25],
    )

    doc.add_heading("第六部分 总体设计", level=1)
    add_para(doc, "系统采用前后端分离、服务层封装、轻量本地存储和外部模型接口解耦的总体架构。前端负责交互与图表展示；后端负责认证、业务编排、数据库访问、向量检索和模型调用；SQLite和Qdrant Local分别承担结构化数据与向量索引存储。")
    add_figure(doc, paths["architecture"], "图6-1 系统架构图")
    add_table(
        doc,
        ["层级", "组件", "设计说明"],
        [
            ["展示层", "Vue3、TypeScript、ECharts", "提供页面导航、对话输入、数据上传、图表与报告展示。"],
            ["接入层", "FastAPI、CORS、Auth Middleware", "统一对外提供REST接口，并在接口入口完成登录检查。"],
            ["业务层", "Analyzer、Dataset Service、Knowledge Service、Auth Service", "封装业务规则、数据处理、知识检索和权限逻辑。"],
            ["数据层", "SQLite、Qdrant Local", "SQLite保存业务数据与会话，Qdrant保存知识向量。"],
            ["模型层", "LLM API、Embedding API", "通过OpenAI兼容协议访问外部模型，支持配置替换。"],
        ],
        [1.05, 2.2, 3.1],
    )

    doc.add_heading("第七部分 用户界面设计", level=1)
    add_para(doc, "界面结构围绕“先登录、再进入工作台、按业务菜单进入功能页面”的思路设计。智能分析页采用对话与报告并列结构，便于展示分析过程和结果。")
    add_figure(doc, paths["ui"], "图7-1 用户界面导航结构图")
    add_table(
        doc,
        ["界面", "主要元素", "交互说明"],
        [
            ["登录页", "账号、密码、登录按钮、错误提示。", "未登录访问系统时首先进入登录页。"],
            ["工作台", "数据集数量、知识片段数量、历史会话、系统状态。", "用于快速查看系统整体运行情况。"],
            ["智能分析", "问题输入框、示例问题、会话过程、图表、查询结果、关键发现、SQL查看。", "支持连续追问和打开旧会话。"],
            ["数据源", "上传CSV、数据集列表、字段预览。", "上传后自动创建数据表和字段元数据。"],
            ["知识库", "知识片段列表、新增表单、删除按钮、向量状态。", "新增/删除后同步影响RAG检索。"],
            ["账户管理", "管理员列表、新增管理员表单。", "只有初始管理员可新增管理员。"],
            ["系统配置", "LLM、Embedding、SQLite、Qdrant状态。", "API Key仅后端可见。"],
        ],
        [1.2, 2.75, 2.4],
    )

    doc.add_heading("第八部分 算法设计", level=1)
    doc.add_heading("一、运行环境说明", level=2)
    add_table(
        doc,
        ["环境项", "说明"],
        [
            ["前端", "Node.js、Vite、Vue 3、TypeScript，默认端口5173。"],
            ["后端", "Python、FastAPI、Uvicorn、Pydantic、httpx，默认端口8000。"],
            ["数据库", "SQLite本地文件backend/storage/data_agent.db。"],
            ["向量库", "Qdrant Local，本地目录backend/storage/qdrant。"],
            ["模型服务", "OpenAI兼容LLM接口和Embedding接口，由backend/.env配置。"],
        ],
        [1.4, 4.9],
    )
    doc.add_heading("二、数据集制作说明", level=2)
    add_para(doc, "系统内置企业经营演示数据，并支持用户上传CSV。数据集制作流程为：读取CSV表头与样例值，识别字段类型，创建SQLite数据表，写入数据集元信息和字段元数据，随后在智能分析时根据元数据识别指标与维度。")
    doc.add_heading("三、核心模块要求", level=2)
    add_figure(doc, paths["algorithm"], "图8-1 智能分析与RAG核心算法流程图")
    add_table(
        doc,
        ["核心模块", "算法/规则", "输出"],
        [
            ["会话上下文合并", "读取上一轮effective_question，追问时继承未被显式替换的时间、指标、维度。", "新的effective_question"],
            ["指标识别", "根据字段名、字段描述和问题关键词识别销售额、订单数、利润、投诉率、转化率等指标。", "metric_sql、metric_label"],
            ["维度识别", "识别地区、产品/展品、品类、渠道等分组维度，支持多维系列图。", "x_field、series_fields"],
            ["SQL生成与校验", "生成聚合SQL后执行只读SQL校验，禁止危险关键字。", "safe_sql或异常提示"],
            ["知识检索", "Embedding问题文本，Qdrant向量检索，并结合关键词加权。", "knowledge_refs"],
            ["回答生成", "数据分析结果使用LLM润色；知识问答要求只依据检索片段回答。", "insights或knowledge_qa answer"],
            ["图表生成", "根据时间趋势、多维分组和结果字段选择line/bar/pie等图表类型。", "chart spec"],
        ],
        [1.4, 3.25, 1.7],
    )

    doc.add_heading("第九部分 运行环境和部署", level=1)
    add_table(
        doc,
        ["节点", "部署内容", "说明"],
        [
            ["浏览器客户端", "访问http://127.0.0.1:5173", "用户登录和使用系统的入口。"],
            ["前端服务", "Vite Dev Server / 静态资源", "开发环境默认端口5173，通过/api代理后端。"],
            ["后端服务", "FastAPI + Uvicorn", "默认端口8000，提供/api/v1接口。"],
            ["SQLite", "本地数据库文件", "保存数据集、元数据、会话、知识片段、管理员和审计日志。"],
            ["Qdrant Local", "本地向量索引目录", "保存知识库向量。"],
            ["外部模型服务", "LLM与Embedding API", "由后端.env配置访问。"],
        ],
        [1.2, 2.2, 2.95],
    )
    add_para(doc, "启动命令：", bold=True, color=NAVY)
    add_para(doc, "后端：cd backend && .venv\\Scripts\\python.exe -m uvicorn app.main:app --host 127.0.0.1 --port 8000", size=10)
    add_para(doc, "前端：cd frontend && npm run dev", size=10)
    add_para(doc, "一键启动：scripts\\start-dev.cmd 或 .\\scripts\\start-dev.ps1", size=10)

    doc.core_properties.title = "数据智能体服务系统概要设计说明书（2/2 图文增强版）"
    doc.core_properties.subject = "概要设计"
    doc.core_properties.author = "数据智能体项目组"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


def main() -> None:
    paths = make_diagrams()
    out = build_doc(paths)
    print(out)
    for name, path in paths.items():
        print(name, path)


if __name__ == "__main__":
    main()
