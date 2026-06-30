from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "数据智能体服务系统_需求规格说明书.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
GRAY = RGBColor(85, 85, 85)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
WHITE = RGBColor(255, 255, 255)


def set_run_font(run, name="Calibri", east_asia="Microsoft YaHei", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_format(paragraph, before=0, after=6, line=1.10, align=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_text(doc, text, style=None, bold=False, color=None, size=11, after=6, before=0):
    p = doc.add_paragraph(style=style)
    set_paragraph_format(p, before=before, after=after, line=1.10)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_format(p, after=4, line=1.167)
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_paragraph_format(p, after=4, line=1.167)
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def table(doc, headers, rows, widths, font_size=9.5):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    set_table_geometry(tbl, widths)
    repeat_table_header(tbl.rows[0])
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_FILL)
        p = cell.paragraphs[0]
        set_paragraph_format(p, after=0, line=1.10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, size=font_size, color=NAVY, bold=True)
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            set_paragraph_format(p, after=0, line=1.10)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if len(str(value)) > 8 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=RGBColor(30, 30, 30))
    doc.add_paragraph()
    return tbl


def add_callout(doc, title, body):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    set_table_geometry(tbl, [9360])
    cell = tbl.rows[0].cells[0]
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    set_paragraph_format(p, after=3, line=1.10)
    r = p.add_run(title)
    set_run_font(r, size=11, color=NAVY, bold=True)
    p2 = cell.add_paragraph()
    set_paragraph_format(p2, after=0, line=1.10)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=RGBColor(60, 60, 60))
    doc.add_paragraph()


def field_run(paragraph, field):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    return run


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.10


def configure_section(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_format(hp, after=0)
    r = hp.add_run("数据智能体服务系统 | 需求规格说明书")
    set_run_font(r, size=9, color=GRAY)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_format(fp, after=0)
    r1 = fp.add_run("第 ")
    set_run_font(r1, size=9, color=GRAY)
    field_run(fp, "PAGE")
    r2 = fp.add_run(" 页")
    set_run_font(r2, size=9, color=GRAY)


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, after=8)
    r = p.add_run("数据智能体服务系统")
    set_run_font(r, size=28, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, after=34)
    r = p.add_run("需求规格说明书")
    set_run_font(r, size=18, color=BLUE, bold=True)

    rows = [
        ("项目周期", "3 周"),
        ("文档版本", "V1.0"),
        ("适用阶段", "需求分析、系统设计、开发实现、测试验收"),
        ("编写日期", "2026 年 6 月 23 日"),
    ]
    table(doc, ["项目", "内容"], rows, [2200, 7160], font_size=11)
    add_callout(
        doc,
        "文档定位",
        "本文档用于说明数据智能体服务系统的建设目标、功能范围、核心需求、验收标准和三周实施计划，可作为项目启动、开发分工和答辩说明的依据。",
    )
    doc.add_page_break()


def add_static_contents(doc):
    doc.add_heading("文档目录", level=1)
    contents = [
        "1. 项目背景与目标",
        "2. 项目范围",
        "3. 用户角色与业务流程",
        "4. 功能需求",
        "5. 非功能需求",
        "6. 数据需求",
        "7. 接口需求",
        "8. 验收标准",
        "9. 三周项目计划",
        "10. 风险与应对",
        "11. 后续扩展方向",
    ]
    for item in contents:
        add_bullet(doc, item)
    doc.add_page_break()


def build_doc():
    doc = Document()
    configure_section(doc)
    configure_styles(doc)
    add_cover(doc)
    add_static_contents(doc)

    doc.add_heading("1. 项目背景与目标", level=1)
    add_text(
        doc,
        "企业在日常经营过程中会产生大量销售、订单、客户投诉、访问转化等数据。传统数据分析系统通常依赖固定报表或人工编写 SQL，存在使用门槛高、分析效率低、指标口径难解释、无法连续追问等问题。",
    )
    add_text(
        doc,
        "本项目建设一个可本地运行的数据智能体服务系统，使用户能够通过自然语言提出经营分析问题，由系统自动完成数据理解、查询规划、安全执行、图表生成、结论总结和知识库问答。",
    )
    add_callout(
        doc,
        "总体目标",
        "完成一个三周周期内可演示、可测试、可扩展的数据智能体 MVP，形成从数据接入到智能分析、知识问答、历史会话和报告导出的完整闭环。",
    )
    table(
        doc,
        ["目标编号", "目标内容"],
        [
            ("G-01", "支持企业经营数据上传、接入、预览和元数据管理。"),
            ("G-02", "支持用户使用自然语言进行数据分析并生成图表。"),
            ("G-03", "支持系统自动生成安全的只读 SQL 查询。"),
            ("G-04", "支持业务知识库问答，用于解释指标口径、字段含义和业务规则。"),
            ("G-05", "支持多轮对话、上下文继承、历史恢复和历史删除。"),
            ("G-06", "支持模型和 Embedding 服务在后端配置，前端不暴露 API Key。"),
        ],
        [1600, 7760],
    )

    doc.add_heading("2. 项目范围", level=1)
    doc.add_heading("2.1 本期建设范围", level=2)
    for item in [
        "前端数据智能体工作台，包括首页、数据源、智能分析、知识库和系统配置页面。",
        "后端 FastAPI 服务，包括数据集、知识库、智能分析、会话和报告接口。",
        "CSV / Excel 数据上传、SQLite 本地存储和数据集元数据识别。",
        "自然语言分析、只读 SQL 安全校验、ECharts 图表展示和关键发现总结。",
        "多轮会话、历史记录、历史恢复和删除能力。",
        "业务知识库新增、删除、Embedding 索引和 Qdrant Local 检索。",
    ]:
        add_bullet(doc, item)
    doc.add_heading("2.2 暂不纳入本期范围", level=2)
    for item in [
        "企业级权限系统，如 OIDC、RBAC、多租户和细粒度行列权限。",
        "真实生产数据库连接器，如 MySQL、PostgreSQL、ClickHouse。",
        "Python 沙箱执行复杂统计脚本。",
        "Word / PDF 高级报告导出和企业模板套打。",
        "大规模分布式部署、模型成本治理和调用限流平台。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("3. 用户角色与业务流程", level=1)
    table(
        doc,
        ["角色", "说明", "核心诉求"],
        [
            ("企业管理者", "查看经营指标和趋势", "快速了解经营变化和异常。"),
            ("运营人员", "分析销售、转化、投诉等业务数据", "通过自然语言快速生成图表。"),
            ("数据分析人员", "验证数据、查看 SQL、解释指标", "保证查询过程透明、可追溯。"),
            ("系统管理员", "配置模型、维护数据和知识库", "管理 API Key、数据源和知识片段。"),
        ],
        [1900, 3300, 4160],
    )
    doc.add_heading("3.1 数据分析流程", level=2)
    for item in [
        "用户输入自然语言问题。",
        "系统识别分析意图，解析指标、时间范围和维度。",
        "系统生成只读 SQL，并经过安全校验。",
        "后端执行查询，返回结构化结果。",
        "前端根据 chart spec 渲染图表，并展示关键发现和明细表格。",
        "系统保存用户问题、助手回答和分析结果，供后续追问和历史恢复使用。",
    ]:
        add_number(doc, item)
    doc.add_heading("3.2 知识库问答流程", level=2)
    for item in [
        "用户提出指标口径、字段含义或业务规则类问题。",
        "系统识别为知识问答意图，不生成 SQL。",
        "系统通过 Embedding 和 Qdrant Local 检索相关知识片段。",
        "系统结合知识片段生成回答，并展示知识依据。",
    ]:
        add_number(doc, item)

    doc.add_heading("4. 功能需求", level=1)
    functional_sections = [
        ("4.1 工作台首页", [
            ("FR-001", "首页概览", "展示数据集数量、知识片段数量、分析次数、会话数量。", "高"),
            ("FR-002", "快捷分析入口", "提供常用分析问题入口，用户点击后进入智能分析。", "中"),
            ("FR-003", "数据资产概览", "展示最近接入的数据集及行数、字段数。", "中"),
        ]),
        ("4.2 数据源管理", [
            ("FR-004", "文件上传", "支持上传 CSV、Excel 文件。", "高"),
            ("FR-005", "数据入库", "上传后自动创建内部数据表并写入数据。", "高"),
            ("FR-006", "元数据识别", "自动识别字段名称、字段类型、样例值、空值率。", "高"),
            ("FR-007", "数据预览", "支持查看数据集字段和部分样例数据。", "中"),
            ("FR-008", "演示数据初始化", "系统首次启动时自动生成企业经营演示数据。", "高"),
        ]),
        ("4.3 智能分析", [
            ("FR-009", "自然语言提问", "用户可输入中文问题进行数据分析。", "高"),
            ("FR-010", "指标识别", "识别销售额、订单数、利润、投诉率、转化率等指标。", "高"),
            ("FR-011", "时间范围识别", "支持近 N 天、近 N 周、本月、本年等时间表达。", "高"),
            ("FR-012", "维度拆分", "支持按地区、产品类别、渠道等维度拆分。", "高"),
            ("FR-013", "多维拆分", "支持日期 + 地区 + 产品类别等多维组合分析。", "高"),
            ("FR-014", "图表类型识别", "支持折线图、柱状图、饼图等展示方式。", "中"),
            ("FR-015", "SQL 生成", "系统根据分析意图生成只读 SQL。", "高"),
            ("FR-016", "查询执行", "系统执行 SQL 并返回结果数据。", "高"),
            ("FR-017", "关键发现生成", "根据查询结果生成业务结论。", "高"),
            ("FR-018", "SQL 查看", "用户可查看系统实际执行的 SQL。", "中"),
        ]),
        ("4.4 图表展示", [
            ("FR-019", "折线图展示", "用于展示时间趋势。", "高"),
            ("FR-020", "柱状图展示", "用于展示分类对比。", "高"),
            ("FR-021", "饼图展示", "用于展示占比类结果。", "中"),
            ("FR-022", "多系列图表", "支持多地区、多产品类别组合系列展示。", "高"),
            ("FR-023", "图例滚动", "多系列过多时支持滚动图例。", "中"),
            ("FR-024", "数据缩放", "时间轴较长时支持横向缩放查看。", "中"),
        ]),
        ("4.5 多轮对话", [
            ("FR-025", "会话创建", "用户首次提问时自动创建会话。", "高"),
            ("FR-026", "上下文继承", "追问时继承上一轮分析条件。", "高"),
            ("FR-027", "条件修改", "支持“只看华东”“改成柱状图”等条件变更。", "高"),
            ("FR-028", "维度追加", "支持“按产品拆分”等追加分析维度。", "高"),
            ("FR-029", "历史消息保存", "保存用户问题、助手回答和分析结果。", "高"),
            ("FR-030", "历史会话恢复", "用户可打开旧会话并查看历史结果。", "高"),
            ("FR-031", "历史会话删除", "用户可删除不需要的历史会话。", "中"),
        ]),
        ("4.6 业务知识库", [
            ("FR-032", "知识片段列表", "展示所有业务知识片段。", "高"),
            ("FR-033", "新增知识片段", "支持新增指标口径、业务规则、字段说明等内容。", "高"),
            ("FR-034", "删除知识片段", "支持删除指定知识片段。", "中"),
            ("FR-035", "向量索引同步", "新增或删除知识后同步更新向量索引。", "高"),
            ("FR-036", "知识库问答", "用户提问指标口径类问题时，系统从知识库检索并回答。", "高"),
            ("FR-037", "知识依据展示", "回答中展示引用的知识片段。", "中"),
        ]),
        ("4.7 系统配置与报告", [
            ("FR-038", "LLM 配置状态", "展示大模型是否已配置。", "高"),
            ("FR-039", "Embedding 配置状态", "展示 Embedding 模型是否已配置。", "高"),
            ("FR-040", "API Key 安全", "API Key 只存储在后端环境变量中，前端不可见。", "高"),
            ("FR-041", "向量库状态", "展示当前向量库类型和已索引数量。", "中"),
            ("FR-042", "HTML 报告导出", "支持将当前会话分析结果导出为 HTML 报告。", "中"),
            ("FR-043", "报告内容", "报告包含结论、数据结果、SQL、知识依据。", "中"),
        ]),
    ]
    for title, rows in functional_sections:
        doc.add_heading(title, level=2)
        table(doc, ["编号", "需求名称", "需求描述", "优先级"], rows, [1200, 1900, 5060, 1200], font_size=9)

    doc.add_heading("5. 非功能需求", level=1)
    nfrs = [
        ("安全性", "前端不得保存或展示模型密钥；SQL 执行限制为只读；危险 SQL 必须拦截；敏感配置不得提交到 Git。"),
        ("易用性", "系统应提供清晰导航，分析结果同时展示图表、结论、表格和 SQL。"),
        ("可维护性", "前后端分离；后端按 routers、services、models 分层；核心分析逻辑与图表展示解耦。"),
        ("可扩展性", "SQLite 可替换为 PostgreSQL 或 MySQL；Qdrant Local 可升级为 Qdrant Server；规划器可升级为 LLM Planner。"),
        ("性能", "常规查询响应应控制在 5 秒以内；上传文件默认限制 20MB；查询结果限制最大返回行数。"),
    ]
    table(doc, ["类别", "要求"], nfrs, [1700, 7660])

    doc.add_heading("6. 数据需求", level=1)
    table(
        doc,
        ["数据对象", "需要保存的主要字段"],
        [
            ("数据集", "名称、描述、来源类型、内部表名、行数、字段数、创建时间、状态。"),
            ("字段元数据", "字段名称、字段类型、字段描述、样例值、空值率。"),
            ("会话", "会话 ID、会话标题、关联数据集、创建时间、更新时间。"),
            ("消息", "角色、内容、分析结果 payload、创建时间。"),
            ("知识片段", "标题、内容、分类、关联数据集、创建时间、向量索引信息。"),
            ("审计日志", "操作类型、资源类型、资源 ID、执行状态、详情和时间。"),
        ],
        [1800, 7560],
    )

    doc.add_heading("7. 接口需求", level=1)
    api_rows = [
        ("GET", "/api/v1/dashboard", "获取首页统计。"),
        ("GET", "/api/v1/config/status", "获取系统配置状态。"),
        ("GET", "/api/v1/datasets", "获取数据集列表。"),
        ("GET", "/api/v1/datasets/{id}", "获取数据集详情。"),
        ("POST", "/api/v1/datasets/upload", "上传数据集。"),
        ("POST", "/api/v1/agent/chat", "智能分析对话。"),
        ("GET", "/api/v1/sessions", "获取历史会话列表。"),
        ("GET", "/api/v1/sessions/{id}", "获取会话详情。"),
        ("DELETE", "/api/v1/sessions/{id}", "删除历史会话。"),
        ("GET", "/api/v1/knowledge", "获取知识片段。"),
        ("POST", "/api/v1/knowledge", "新增知识片段。"),
        ("DELETE", "/api/v1/knowledge/{id}", "删除知识片段。"),
        ("POST", "/api/v1/knowledge/reindex", "重建知识库向量索引。"),
        ("GET", "/api/v1/knowledge/vector/status", "获取向量索引状态。"),
        ("GET", "/api/v1/reports/{session_id}.html", "导出 HTML 报告。"),
    ]
    table(doc, ["方法", "接口", "用途"], api_rows, [1200, 3800, 4360], font_size=9)

    doc.add_heading("8. 验收标准", level=1)
    accept_rows = [
        ("数据上传", "能成功上传 CSV 文件并生成数据集。"),
        ("元数据展示", "能展示字段、类型、样例值。"),
        ("自然语言分析", "输入经营问题后能返回图表和结论。"),
        ("SQL 安全", "危险 SQL 被拦截。"),
        ("多轮追问", "第二轮问题能继承上一轮条件。"),
        ("多维图表", "地区 + 产品类别等组合维度能完整展示。"),
        ("知识库问答", "能回答指标口径类问题并展示知识依据。"),
        ("历史会话", "能查看、恢复、删除历史对话。"),
        ("知识片段", "能新增和删除知识片段。"),
        ("配置安全", "前端不展示 API Key 明文。"),
    ]
    table(doc, ["验收项", "验收标准"], accept_rows, [2200, 7160])
    add_callout(
        doc,
        "测试验收要求",
        "后端核心单元测试应全部通过；前端生产构建应成功；智能分析、知识库问答、删除历史会话和删除知识片段等核心链路应能完成端到端验证。",
    )

    doc.add_heading("9. 三周项目计划", level=1)
    plan_rows = [
        ("第 1 天", "完成需求分析和需求文档", "需求规格说明书"),
        ("第 2 天", "完成系统架构设计", "架构设计文档"),
        ("第 3 天", "搭建前后端项目骨架", "FastAPI + Vue 项目"),
        ("第 4 天", "设计数据库表结构", "SQLite 表结构"),
        ("第 5 天", "实现演示数据初始化和首页概览", "首页 MVP"),
        ("第 6 天", "实现 CSV / Excel 上传", "数据源管理模块"),
        ("第 7 天", "实现元数据识别和数据预览", "数据集详情页面"),
        ("第 8 天", "实现自然语言分析规划器", "分析服务"),
        ("第 9 天", "实现 SQL 安全校验和查询执行", "安全查询模块"),
        ("第 10 天", "实现图表展示和关键发现", "智能分析页面"),
        ("第 11 天", "实现多轮对话和历史记录", "会话管理"),
        ("第 12 天", "实现知识库新增、删除、检索", "知识库模块"),
        ("第 13 天", "接入 Embedding 和 Qdrant Local", "RAG 检索能力"),
        ("第 14 天", "完善多维图表、报告导出、配置页面", "完整演示系统"),
        ("第 15 天", "测试、修复、整理答辩材料", "测试报告、答辩材料"),
    ]
    table(doc, ["时间", "工作内容", "交付物"], plan_rows, [1300, 5100, 2960], font_size=9)

    doc.add_heading("10. 风险与应对", level=1)
    risk_rows = [
        ("大模型 API 不稳定", "分析结论生成失败", "保留本地规则分析和降级逻辑。"),
        ("Embedding 服务不可用", "知识库语义检索失败", "降级为关键词检索。"),
        ("SQL 生成不准确", "查询结果错误", "使用确定性规则规划器，并展示 SQL 供用户检查。"),
        ("图表维度过多", "前端展示混乱", "支持滚动图例、缩放条、多字段系列组合。"),
        ("三周周期较短", "功能范围过大", "优先完成 MVP，企业级功能列入后续计划。"),
        ("API Key 泄露", "安全风险", "API Key 只存后端 .env，前端不返回密钥。"),
    ]
    table(doc, ["风险", "影响", "应对措施"], risk_rows, [2200, 2600, 4560], font_size=9)

    doc.add_heading("11. 后续扩展方向", level=1)
    for item in [
        "接入 MySQL、PostgreSQL、ClickHouse 等企业数据库。",
        "增加权限系统，支持用户、角色、数据集权限。",
        "增加 Python 沙箱，用于复杂统计分析。",
        "增加 Word / PDF 报告导出。",
        "增加模型调用监控、成本统计和限流。",
        "建设离线评测集，持续评估智能体回答质量。",
        "增加数据血缘、指标血缘和异常归因能力。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("12. 总结", level=1)
    add_text(
        doc,
        "本项目旨在建设一个面向企业经营数据分析的数据智能体服务系统。系统以自然语言交互为入口，结合结构化数据查询、业务知识库检索、多轮会话管理和图表报告展示，实现从用户提问到分析结论的完整闭环。",
    )
    add_text(
        doc,
        "本期三周项目重点完成可演示、可运行、可扩展的 MVP，为后续企业级数据智能体平台建设奠定基础。",
    )
    doc.core_properties.title = "数据智能体服务系统需求规格说明书"
    doc.core_properties.subject = "需求文档"
    doc.core_properties.author = "DataAgent 项目组"
    doc.core_properties.keywords = "数据智能体,需求规格说明书,自然语言分析,RAG,多轮对话"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
