from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path.home() / "Downloads" / "《XXXXX系统》项目需求说明-模版.docx"
OUT = ROOT / "docs" / "数据智能体服务系统_项目需求说明_模板版.docx"


def set_run_font(run, name="宋体", east_asia="宋体", size=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def para_text(paragraph, text, size=None, bold=None, align=None):
    paragraph.clear()
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return paragraph


def clear_cell(cell):
    for p in cell.paragraphs:
        p.clear()
    for table in list(cell.tables):
        cell._tc.remove(table._tbl)


def cell_text(cell, text, size=10.5, bold=False, align=None):
    clear_cell(cell)
    parts = str(text).split("\n")
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    for idx, part in enumerate(parts):
        if idx:
            p.add_run().add_break()
        run = p.add_run(part)
        set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_shading(cell, fill="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def clone_row(table):
    tr = deepcopy(table.rows[-1]._tr)
    table._tbl.append(tr)
    return table.rows[-1]


def ensure_rows(table, count):
    while len(table.rows) < count:
        clone_row(table)
    while len(table.rows) > count:
        table._tbl.remove(table.rows[-1]._tr)


def fill_table(table, rows, header=True):
    ensure_rows(table, len(rows))
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell_text(cell, value, size=9.5 if len(str(value)) > 20 else 10.5, bold=header and r_idx == 0)
            if header and r_idx == 0:
                set_cell_shading(cell, "D9EAF7")


def replace_first_matching(paragraphs, old, new):
    for p in paragraphs:
        if old in p.text:
            para_text(p, new)
            return True
    return False


def fill_single_cell_table(table, text):
    cell_text(table.rows[0].cells[0], text, size=10.5)


def build():
    if not TEMPLATE.exists():
        raise FileNotFoundError(TEMPLATE)
    doc = Document(TEMPLATE)

    # Cover
    para_text(doc.paragraphs[4], "《数据智能体服务系统》", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para_text(doc.paragraphs[5], "用 户 需 求 说 明 书", size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para_text(doc.paragraphs[9], "v1.0", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    para_text(doc.paragraphs[12], "数据智能体项目组", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Version history
    fill_table(
        doc.tables[1],
        [
            ["版本/状态", "作者", "参与者", "日期", "备注"],
            ["V1.0 / 初稿", "项目组", "产品、前端、后端、测试", "2026-06-23", "完成项目需求说明书模板版"],
            ["V1.1 / 评审稿", "项目组", "指导老师、项目成员", "待定", "根据评审意见修订"],
            ["V1.2 / 定稿", "项目组", "项目成员", "待定", "作为开发与验收依据"],
        ],
    )

    # Static contents
    toc = [
        "第一部分 引言\t3",
        "一、说明\t3",
        "二、定义\t4",
        "第二部分 综述\t4",
        "一、项目背景\t4",
        "二、建设目标\t4",
        "三、建设原则\t5",
        "四、用户业务需求说明\t6",
        "1、整体业务需求示意图\t6",
        "2、需求详细说明\t6",
        "2.1 数据分析业务流程描述\t6",
        "1）、自然语言数据分析流程\t6",
        "2）、知识库问答流程\t6",
        "第三部分 需求分析\t7",
        "一、用例分析\t7",
        "1、智能分析用例\t7",
        "1.1、智能分析用例描述\t7",
        "二、界面风格\t7",
        "第四部分 验收标准\t8",
        "一、功能范围定义\t8",
        "二、性能指标定义\t9",
        "第五部分 环境和部署要求\t11",
        "一、网络部署图\t11",
        "二、应用部署图\t11",
        "三、其他部署\t11",
        "四、运行环境说明\t12",
        "1、服务器\t12",
        "2、客户机器环境\t12",
        "3、模型与向量库环境\t12",
    ]
    for idx, line in enumerate(toc, start=21):
        para_text(doc.paragraphs[idx], line)

    # Headings and sections
    fill_single_cell_table(
        doc.tables[2],
        "本文档用于说明“数据智能体服务系统”的用户需求、业务流程、功能范围、性能要求、验收标准和部署环境要求。\n"
        "文档读者包括项目负责人、开发人员、测试人员、指导老师及后续维护人员。\n"
        "本文档经评审确认后，将作为系统设计、开发实现、测试验收和答辩说明的依据。",
    )
    fill_single_cell_table(
        doc.tables[3],
        "数据智能体：能够根据用户自然语言问题自动完成数据理解、查询规划、分析执行和结果解释的软件服务。\n"
        "RAG：检索增强生成，通过知识库检索补充模型回答依据。\n"
        "Embedding：将文本转换为向量，用于语义检索。\n"
        "Qdrant Local：本项目采用的本地轻量级向量数据库。\n"
        "只读 SQL：仅允许 SELECT/WITH 查询，不允许修改数据的 SQL 语句。\n"
        "多轮会话：系统保存上下文，支持用户基于上一轮结果继续追问。",
    )
    fill_single_cell_table(
        doc.tables[4],
        "企业在经营过程中会产生大量销售、订单、客户投诉、访问转化等数据。传统数据看板通常只能展示预设图表，用户如果想临时更换指标、维度或继续追问，需要人工编写 SQL 或依赖数据分析人员。\n"
        "本项目面向企业经营数据分析场景，建设一个支持自然语言交互的数据智能体系统，使用户能够直接提出经营分析问题，系统自动完成数据查询、图表展示、关键发现总结和指标口径解释。\n"
        "系统可提升业务人员自主分析能力，降低数据分析门槛，并为后续企业级智能数据分析平台建设提供基础。",
    )
    fill_single_cell_table(
        doc.tables[5],
        "1、降低数据分析门槛，用户无需编写 SQL 即可完成经营分析。\n"
        "2、提高分析效率，支持自然语言提问、自动生成图表和关键结论。\n"
        "3、规范指标口径，通过知识库统一解释销售额、投诉率、转化率等指标。\n"
        "4、支持多轮追问，能够继承时间范围、指标、地区、产品等上下文条件。\n"
        "5、保证执行安全，通过只读 SQL 校验限制危险操作。\n"
        "6、形成可演示、可扩展、可部署的三周项目 MVP。",
    )

    principle_texts = [
        "系统功能以真实可用为优先，围绕数据上传、智能分析、知识问答、历史会话和图表展示形成完整闭环，避免只做静态页面或演示假数据。",
        "系统采用前后端分离、模块化接口设计，并预留模型、向量库和数据源替换能力，便于后续从 SQLite 扩展到企业级数据库。",
        "前端界面应简洁清晰，分析结果应同时提供图表、表格、关键发现和 SQL 查看入口，便于用户理解结果来源。",
        "系统支持 CSV/Excel 数据接入，后续可扩展 MySQL、PostgreSQL、ClickHouse 等数据库；Qdrant Local 后续可平滑升级为 Qdrant Server。",
        "模型 API Key 和 Embedding API Key 只存储在后端环境变量中；SQL 执行必须经过只读校验，防止误删、越权和危险操作。",
    ]
    for idx, text in enumerate(principle_texts, start=6):
        fill_single_cell_table(doc.tables[idx], text)

    # Business process
    fill_single_cell_table(
        doc.tables[11],
        "用户自然语言提问\n→ 智能体识别意图\n→ 解析指标、时间、维度\n→ 生成只读 SQL\n→ SQL 安全校验\n→ 执行查询\n→ 生成图表、表格和关键发现\n→ 保存会话历史\n\n"
        "知识库问答链路：\n用户提出指标口径问题 → Embedding 向量化 → Qdrant Local 检索知识片段 → 生成回答 → 展示知识依据。",
    )
    para_text(doc.paragraphs[92], "2.1\t数据分析业务流程描述")
    para_text(doc.paragraphs[93], "1）、自然语言数据分析流程")
    fill_single_cell_table(
        doc.tables[12],
        "用户进入“智能分析”页面后，选择数据集并输入分析问题，例如“分析近10天各地区销售额趋势”。系统识别问题中的时间范围、指标和维度，生成只读 SQL 并执行查询。查询结果返回后，前端使用 ECharts 展示折线图或柱状图，同时展示关键发现和明细数据。\n"
        "用户可继续追问，例如“按展品拆分”或“只看华东，并改成柱状图”。系统读取上一轮 effective_question，继承上下文并生成新的分析结果。",
    )
    para_text(doc.paragraphs[96], "2）、知识库问答流程")
    fill_single_cell_table(
        doc.tables[13],
        "用户进入“智能分析”页面后，可提出指标口径或业务规则类问题，例如“投诉率的计算口径是什么”。系统识别为知识库问答，不生成 SQL，而是从业务知识库中检索相关片段。\n"
        "知识片段通过 Embedding 转换为向量并存储在 Qdrant Local 中。检索结果会作为回答依据返回前端，用户可以看到系统引用的知识片段。",
    )

    # Use case analysis
    para_text(doc.paragraphs[106], "1、智能分析用例")
    fill_single_cell_table(
        doc.tables[14],
        "核心用例包括：数据集上传、自然语言分析、多轮追问、知识库问答、历史会话恢复、历史会话删除、知识片段新增、知识片段删除、HTML 报告导出和系统配置查看。",
    )
    para_text(doc.paragraphs[108], "智能分析用例说明：")
    para_text(doc.paragraphs[109], "1.1、智能分析用例描述")
    use_case_rows = [
        ["ID", "UC-001"],
        ["用例名称", "自然语言经营数据分析"],
        ["父用例ID", "无"],
        ["主要执行者", "企业管理者、运营人员、数据分析人员"],
        ["前置条件", "系统已启动；至少存在一个可用数据集；用户已进入智能分析页面。"],
        ["事件流", "1、用户输入自然语言问题。\n2、系统识别指标、时间范围和分析维度。\n3、系统生成只读 SQL 并进行安全校验。\n4、系统执行查询并返回结果。\n5、前端展示图表、关键发现、数据表格和 SQL。"],
        ["可选事件流", "用户可继续追问，如“只看华东”“按产品类别拆分”“改成柱状图”；系统继承上一轮上下文并重新分析。"],
        ["异常事件流", "若数据集不存在，提示用户先上传数据；若 SQL 校验失败，拒绝执行；若模型或向量服务不可用，进入本地规则或关键词检索降级流程。"],
        ["后置条件", "系统保存用户问题、助手回答和分析结果，用户可在历史会话中恢复查看。"],
    ]
    fill_table(doc.tables[15], use_case_rows, header=False)

    # UI descriptions
    fill_single_cell_table(
        doc.tables[16],
        "系统登录/入口页面采用简洁管理后台风格。左侧为功能导航，顶部展示当前页面标题和数据集选择器。系统当前 MVP 面向本地演示环境，可直接进入工作台；后续可扩展账号登录与权限控制。",
    )
    fill_single_cell_table(
        doc.tables[17],
        "主页面展示数据集数量、知识片段数量、累计智能分析次数和分析会话数量。同时提供常用分析场景快捷入口和最近数据资产列表，方便用户快速进入分析。",
    )
    para_text(doc.paragraphs[118], "智能分析功能页面：")
    fill_single_cell_table(
        doc.tables[18],
        "智能分析页面左侧为对话面板，包括历史对话、消息列表、快捷提问和输入框；右侧为分析结果区域，包括图表、关键发现、知识依据、SQL 和明细数据表。图表支持多系列展示，例如“区域 / 产品类别”组合系列。",
    )

    # Acceptance - functional scope
    functional_rows = [
        ["#", "产品", "模块", "组件", "规格/型号", "角色", "接入"],
        ["1", "数据智能体服务系统", "工作台首页", "统计概览", "数据集、知识片段、分析次数、会话数量", "所有用户", "Web"],
        ["2", "数据智能体服务系统", "数据源管理", "文件上传", "CSV、Excel 上传与入库", "系统管理员、分析人员", "Web/API"],
        ["3", "数据智能体服务系统", "数据源管理", "元数据识别", "字段类型、样例值、空值率", "分析人员", "Web/API"],
        ["4", "数据智能体服务系统", "智能分析", "自然语言提问", "中文经营问题输入", "管理者、运营人员", "Web/API"],
        ["5", "数据智能体服务系统", "智能分析", "SQL 规划", "指标、时间、维度解析并生成只读 SQL", "系统", "API"],
        ["6", "数据智能体服务系统", "智能分析", "图表展示", "折线图、柱状图、饼图、多系列图表", "所有用户", "Web"],
        ["7", "数据智能体服务系统", "多轮会话", "上下文继承", "继承时间、指标、地区、产品等条件", "所有用户", "Web/API"],
        ["8", "数据智能体服务系统", "会话管理", "历史会话", "查看、恢复、删除历史对话", "所有用户", "Web/API"],
        ["9", "数据智能体服务系统", "知识库", "知识片段管理", "新增、查看、删除业务知识片段", "系统管理员", "Web/API"],
        ["10", "数据智能体服务系统", "知识库", "RAG 问答", "Embedding + Qdrant Local 检索", "所有用户", "Web/API"],
        ["11", "数据智能体服务系统", "系统配置", "模型配置状态", "LLM、Embedding、向量库状态展示", "系统管理员", "Web"],
        ["12", "数据智能体服务系统", "报告导出", "HTML 报告", "导出结论、数据、SQL、知识依据", "分析人员", "Web/API"],
    ]
    fill_table(doc.tables[19], functional_rows)

    performance_rows = [
        ["#", "产品", "模块", "组件", "规格/型号", "性能级别"],
        ["1", "数据智能体服务系统", "智能分析", "常规查询响应", "演示数据集下 5 秒内返回图表和结论", "A"],
        ["2", "数据智能体服务系统", "数据上传", "文件大小", "默认支持 20MB 以内 CSV / Excel", "B"],
        ["3", "数据智能体服务系统", "SQL 安全", "危险 SQL 拦截", "DELETE、UPDATE、DROP 等语句必须拒绝", "A"],
        ["4", "数据智能体服务系统", "会话管理", "历史保存", "用户消息、助手回答和 payload 持久化", "A"],
        ["5", "数据智能体服务系统", "知识库", "向量检索", "Embedding 配置可用时使用 Qdrant Local 检索", "A"],
        ["6", "数据智能体服务系统", "知识库", "降级检索", "Embedding 不可用时降级关键词检索", "B"],
        ["7", "数据智能体服务系统", "前端展示", "多系列图表", "支持地区 + 产品类别等组合系列显示", "A"],
        ["8", "数据智能体服务系统", "配置安全", "API Key", "仅由后端 .env 读取，前端不返回明文", "A"],
        ["9", "数据智能体服务系统", "可维护性", "代码结构", "前后端分离，后端 routers/services/models 分层", "B"],
    ]
    fill_table(doc.tables[20], performance_rows)

    # Deployment
    fill_single_cell_table(
        doc.tables[21],
        "浏览器客户端\n→ http://127.0.0.1:5173 前端 Vite 服务\n→ /api/v1 反向代理或直连 FastAPI\n→ SQLite 本地数据文件\n→ Qdrant Local 向量索引目录\n→ 外部 LLM / Embedding API 服务",
    )
    fill_single_cell_table(
        doc.tables[22],
        "前端应用：Vue 3 + TypeScript + Vite + ECharts。\n"
        "后端应用：FastAPI + Pydantic + Uvicorn。\n"
        "结构化存储：SQLite，保存数据集、字段、会话、消息、知识片段和审计日志。\n"
        "向量存储：Qdrant Local，保存业务知识片段向量。\n"
        "模型服务：通过 OpenAI 兼容接口接入 DeepSeek、SiliconFlow 等模型服务。",
    )
    para_text(doc.paragraphs[168], "三、其他部署")
    fill_single_cell_table(
        doc.tables[23],
        "本期采用本地轻量化部署方案，不强制部署独立数据库服务器或向量数据库服务器。后续生产环境可扩展为 Nginx + FastAPI + PostgreSQL + Qdrant Server + 对象存储的部署架构。",
    )
    fill_single_cell_table(
        doc.tables[24],
        "开发/演示服务器建议配置：Windows 10/11，Python 3.11+，Node.js 18+，内存 8GB 以上，磁盘剩余空间 5GB 以上。后端端口默认 8000，前端端口默认 5173。",
    )
    fill_single_cell_table(
        doc.tables[25],
        "客户机器环境：Microsoft Edge、Chrome 或其他现代浏览器；建议分辨率 1366×768 以上；可访问本机或局域网部署的前端地址。",
    )
    para_text(doc.paragraphs[177], "3、模型与向量库环境")
    fill_single_cell_table(
        doc.tables[26],
        "后端 .env 配置 LLM_API_KEY、LLM_BASE_URL、LLM_MODEL、EMBEDDING_API_KEY、EMBEDDING_BASE_URL、EMBEDDING_MODEL、VECTOR_STORE、QDRANT_PATH 等参数。\n"
        "API Key 只由后端读取，前端仅展示配置状态，不展示密钥明文。Qdrant Local 默认存储路径为 backend/storage/qdrant。",
    )

    doc.core_properties.title = "数据智能体服务系统项目需求说明"
    doc.core_properties.subject = "项目需求说明"
    doc.core_properties.author = "数据智能体项目组"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
